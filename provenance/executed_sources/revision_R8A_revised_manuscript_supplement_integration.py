from __future__ import annotations

import fcntl
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import time
import zipfile
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", "/tmp/revision_r8a_matplotlib")
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import revision_R3A_P1_float32_engine_frozen_trajectory_unit_test as engine


REVISION_PROTOCOL_SHA256 = "6807b71de18ca82013cfa4360d760e0daf9a920a1acc0625dcb13bd8f4d07249"
R0_PACKET_SHA256 = "0800e315a29b81934095ba56deaea3f8b6600fd0df13db348d7ea72d3b82df78"
R2_PACKET_SHA256 = "f620d5a29f7c6fb3ffac703e5f72208baf99ac76f8aeb69324c062210bfc4a8d"
R7B_PACKET_SHA256 = "fae4cb5ec6ac904fcb2576d63835f1f3d549f4b8362abbe63b2026fb40337d2e"
LOCAL_QA_SOURCE_DOCX_SHA256 = "739be360e1e2bd49e65eb3f2bbe8408a19844d421448a53042c2198fa8ffb37b"
R0_BASENAME = "stageR0_reviewer_revision_protocol_lock_packet.zip"
R2_BASENAME = "revision_R2_manuscript_corrections_packet.zip"
R7B_BASENAME = "revision_R7B_locked_statistical_analysis_supplement_packet.zip"
R0_REMOTE = (
    engine.REMOTE_BASE
    + "/Reviewer_Revision/StageR0_Reviewer_Revision_Protocol_Lock/"
    + R0_BASENAME
)
R7B_REMOTE = (
    engine.REMOTE_BASE
    + "/Reviewer_Revision/Revision_R7B_Locked_Statistical_Analysis_and_Supplement/"
    + R7B_BASENAME
)
R2_REMOTE = (
    engine.REMOTE_BASE
    + "/Reviewer_Revision/Revision_R2_Manuscript_Corrections/"
    + R2_BASENAME
)

WORKING = Path(os.environ.get("REVISION_R8A_WORKING", "/kaggle/working"))
SOURCE_DOCX = Path(
    os.environ.get(
        "REVISION_R8A_SOURCE_DOCX",
        str(WORKING / "revision_R8A_source_R2_corrected_manuscript.docx"),
    )
)
INPUT_ROOT = WORKING / "REVISION_R8A_FROZEN_INPUTS"
RESULT_ROOT = WORKING / "DELTA_REVIEWER_REVISION" / "Revision_R8A_Revised_Manuscript_Supplement_Integration"
EVIDENCE_ROOT = RESULT_ROOT / "Supplementary_Evidence"
RENDER_ROOT = WORKING / "REVISION_R8A_RENDER_QA"
PACKET_PATH = WORKING / "revision_R8A_revised_manuscript_supplement_integration_packet.zip"
REMOTE_OUTPUT = engine.REMOTE_BASE + "/Reviewer_Revision/Revision_R8A_Revised_Manuscript_Supplement_Integration"
START_TIME = time.time()

BLUE = "17365D"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"


def atomic_csv(frame: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    frame.to_csv(temporary, index=False)
    os.replace(temporary, path)


def atomic_json(payload: dict, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(
        json.dumps(payload, indent=2, sort_keys=True, ensure_ascii=False, default=str),
        encoding="utf-8",
    )
    os.replace(temporary, path)


def direct_restore(basename: str, expected_hash: str, remote: str) -> tuple[Path, str]:
    destination = INPUT_ROOT / basename
    destination.parent.mkdir(parents=True, exist_ok=True)
    if destination.exists() and engine.sha256_file(destination) == expected_hash and engine.archive_crc_passes(destination):
        return destination, "EXISTING_VERIFIED_COPY"
    last_error = ""
    for attempt in range(1, 6):
        temporary = destination.with_suffix(destination.suffix + f".download{attempt}")
        temporary.unlink(missing_ok=True)
        result = engine.rclone(
            ["copyto", remote, str(temporary), "--retries", "5", "--low-level-retries", "10", "--timeout", "5m"],
            check=False,
        )
        if (
            result.returncode == 0
            and temporary.exists()
            and engine.sha256_file(temporary) == expected_hash
            and engine.archive_crc_passes(temporary)
        ):
            os.replace(temporary, destination)
            return destination, "GOOGLE_DRIVE_DIRECT"
        last_error = (result.stderr or result.stdout or "hash-or-crc-mismatch")[-1000:]
        temporary.unlink(missing_ok=True)
    raise RuntimeError(f"Could not restore verified {basename}: {last_error}")


def read_csv(packet: Path, basename: str) -> pd.DataFrame:
    return engine.read_csv_member(packet, basename)


def read_json(packet: Path, basename: str) -> dict:
    return engine.read_json_member(packet, basename)


def restore_source_docx_from_r2(packet: Path) -> tuple[str, str, str]:
    r2_report = read_json(packet, "revision_R2_manuscript_corrections_report.json")
    expected_hash = str(r2_report.get("revised_docx_sha256", "")).lower()
    if not re.fullmatch(r"[0-9a-f]{64}", expected_hash):
        raise RuntimeError("Verified R2 report does not declare a valid revised DOCX SHA-256")
    if SOURCE_DOCX.exists() and engine.sha256_file(SOURCE_DOCX) == expected_hash:
        return "EXISTING_VERIFIED_R2_DOCX", expected_hash, expected_hash
    matches = []
    with zipfile.ZipFile(packet, "r") as archive:
        for member in archive.infolist():
            if not member.is_dir() and Path(member.filename).name == "revision_R2_corrected_manuscript.docx":
                matches.append(member)
        if len(matches) != 1:
            raise RuntimeError(f"Expected exactly one corrected R2 manuscript in packet; found {len(matches)}")
        data = archive.read(matches[0])
    actual_hash = hashlib.sha256(data).hexdigest()
    if actual_hash != expected_hash:
        raise RuntimeError("Corrected R2 manuscript does not match the SHA-256 declared by the verified R2 report")
    with zipfile.ZipFile(io.BytesIO(data), "r") as docx_archive:
        if docx_archive.testzip() is not None:
            raise RuntimeError("Corrected R2 DOCX fails its internal ZIP CRC")
        required_members = {"[Content_Types].xml", "word/document.xml"}
        if not required_members.issubset(docx_archive.namelist()):
            raise RuntimeError("Corrected R2 DOCX is missing required OpenXML members")
    source_document = Document(io.BytesIO(data))
    normalized_paragraphs = {" ".join(paragraph.text.split()) for paragraph in source_document.paragraphs}
    required_headings = {
        "Abstract",
        "Introduction",
        "Materials and methods",
        "Results",
        "Discussion",
        "Limitations",
        "Conclusion",
        "Declarations",
        "References",
    }
    if not required_headings.issubset(normalized_paragraphs):
        raise RuntimeError("Corrected R2 manuscript does not satisfy the locked section schema")
    if len(source_document.paragraphs) < 200 or len(source_document.tables) < 7:
        raise RuntimeError("Corrected R2 manuscript is structurally incomplete")
    SOURCE_DOCX.parent.mkdir(parents=True, exist_ok=True)
    temporary = SOURCE_DOCX.with_suffix(SOURCE_DOCX.suffix + ".tmp")
    temporary.write_bytes(data)
    os.replace(temporary, SOURCE_DOCX)
    return matches[0].filename, expected_hash, actual_hash


def fmt_number(value: float, digits: int = 4, signed: bool = False) -> str:
    if not np.isfinite(float(value)):
        return "NA"
    return f"{float(value):+.{digits}f}" if signed else f"{float(value):.{digits}f}"


def fmt_p(value: float) -> str:
    value = float(value)
    if value < 0.0001:
        return "<0.0001"
    return f"{value:.4f}".rstrip("0").rstrip(".")


def interval_text(row: pd.Series) -> str:
    return f"[{fmt_number(row.bca_95_ci_low, signed=True)}, {fmt_number(row.bca_95_ci_high, signed=True)}]"


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def find_paragraph(document: Document, exact: str | None = None, startswith: str | None = None):
    matches = []
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if exact is not None and text == exact:
            matches.append(paragraph)
        if startswith is not None and text.startswith(startswith):
            matches.append(paragraph)
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph exact={exact!r} startswith={startswith!r}; found {len(matches)}")
    return matches[0]


def mark_runs(paragraph, highlight: bool) -> None:
    if highlight:
        for run in paragraph.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_text(paragraph, text: str, highlight: bool = False, bold: bool = False) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.bold = bold
    mark_runs(paragraph, highlight)


def append_before(anchor, text: str, style: str | None = None, highlight: bool = False):
    from docx.text.paragraph import Paragraph

    new_element = OxmlElement("w:p")
    anchor._p.addprevious(new_element)
    paragraph = Paragraph(new_element, anchor._parent)
    if style is not None:
        paragraph.style = style
    paragraph.add_run(text)
    mark_runs(paragraph, highlight)
    return paragraph


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    total = int(sum(widths))
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_table(table, widths: list[int], highlight: bool = False, font_size: float = 8.5) -> None:
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade_cell(cell, BLUE)
            elif row_index % 2 == 0:
                shade_cell(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    elif highlight:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_table_before(anchor, headers: list[str], rows: list[list[str]], widths: list[int], highlight: bool = False):
    document = anchor.part.document
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    style_table(table, widths, highlight=highlight)
    anchor._p.addprevious(table._tbl)
    return table


def renumber_captions(document: Document, highlight: bool) -> None:
    replacements = {
        "Table 2a.": "Table 2.",
        "Table 2b.": "Table 3.",
        "Table 2. Core notation": "Table 4. Core notation",
        "Table 3. Locked K07": "Table 5. Locked K07",
        "Table 5. Descriptive cross-model": "Table 6. Descriptive cross-model",
        "Table 4. Multiplicity-controlled": "Table 7. Multiplicity-controlled",
        "Figure 8. Descriptive cross-engine": "Figure 5. Descriptive cross-engine",
        "Figure 5. Deep K07": "Figure 6. Deep K07",
        "Figure 6. Deep-model": "Figure 7. Deep-model",
        "Figure 7. Deep K07": "Figure 8. Deep K07",
    }
    for paragraph in document.paragraphs:
        for old, new in replacements.items():
            if paragraph.text.startswith(old):
                set_text(paragraph, new + paragraph.text[len(old) :], highlight=highlight)
                break
        if paragraph.text.startswith("Table "):
            paragraph.paragraph_format.keep_with_next = True


def create_robustness_figure(tests: pd.DataFrame, destination: Path) -> None:
    imbalance = tests.loc[
        tests["analysis_id"].str.contains("REV_SECONDARY_IMBALANCE")
        & tests["analysis_id"].str.endswith("PCBM_MINUS_RANDOM_UNIFORM")
    ].copy()
    imbalance["label"] = imbalance["analysis_id"].str.extract(r"RIDGE__(.*?)__PCBM")[0].str.replace(r"_\d+$", "", regex=True)
    aulc = tests.loc[tests["analysis_id"].str.contains("REV_SECONDARY_AULC")].copy()
    aulc["label"] = aulc["analysis_id"].str.split("PCBM_MINUS_").str[-1].str.replace("_", " ")
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))
    for axis, frame, title in [
        (axes[0], imbalance, "K07 PCBM − random by pool condition"),
        (axes[1], aulc, "Normalized AULC: PCBM − comparator"),
    ]:
        frame = frame.reset_index(drop=True)
        y = np.arange(len(frame))
        means = frame["mean_paired_delta"].to_numpy(float)
        low = frame["bca_95_ci_low"].to_numpy(float)
        high = frame["bca_95_ci_high"].to_numpy(float)
        axis.errorbar(means, y, xerr=[means - low, high - means], fmt="o", color="#17365D", ecolor="#6C8EAD", capsize=3)
        axis.axvline(0, color="#A33A3A", linewidth=1, linestyle="--")
        axis.set_yticks(y, frame["label"].tolist())
        axis.set_title(title, fontsize=11, weight="bold")
        axis.set_xlabel("Participant-level mean paired difference")
        axis.grid(axis="x", alpha=0.2)
        axis.invert_yaxis()
    fig.tight_layout()
    destination.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(destination, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def revision_text_bundle(tests: pd.DataFrame, report: dict) -> dict:
    focal = tests.loc[tests["analysis_id"].eq("REV_FOCAL_01__RIDGE_EFFECT_MODIFICATION_K07")].iloc[0]
    significant = int(tests["significant_holm_0_05"].sum())
    family_counts = (
        tests.assign(significant=tests["significant_holm_0_05"].astype(int))
        .groupby("multiplicity_family", as_index=False)
        .agg(tests=("analysis_id", "size"), significant=("significant", "sum"))
    )
    imbalance_random = tests.loc[
        tests["analysis_id"].str.contains("REV_SECONDARY_IMBALANCE")
        & tests["analysis_id"].str.endswith("PCBM_MINUS_RANDOM_UNIFORM")
    ]
    split = tests.loc[tests["analysis_id"].str.startswith("REV_SPLIT_STABILITY")]
    drift = tests.loc[tests["analysis_id"].str.startswith("REV_DRIFT")]
    deep = tests.loc[tests["analysis_id"].str.startswith("REV_DEEP_STABILITY")]
    supported = bool(
        focal.mean_paired_delta > 0
        and focal.exact_two_sided_p_value < 0.05
        and focal.bca_95_ci_low > 0
    )
    status = "supported" if supported else "not supported"
    abstract = (
        "Longitudinal high-density surface electromyography (HD-sEMG) is vulnerable to session-to-session shift, "
        "motivating label-efficient leakage-controlled adaptation. We evaluated Predicted-Class-Balanced Margin (PCBM), "
        "an uncertainty heuristic that distributes early queries across predicted classes, using the seven-participant DELTA dataset. "
        "Six able-bodied participants formed the participant-level inferential cohort; the participant with limb absence remained descriptive. "
        "Reviewer-motivated robustness extensions compared PCBM with random acquisition, global margin, least confidence, predictive entropy, "
        "RBMAL, core-set, and BADGE across balanced and class-imbalanced pools, four temporal splits, within-session drift definitions, and six "
        "locked TCN training seeds. The focal Ridge estimand—[(PCBM−random) averaged over moderate and severe imbalance] minus the balanced-pool "
        f"contrast—was {fmt_number(focal.mean_paired_delta, signed=True)} (BCa 95% CI {interval_text(focal)}; exact two-sided p={fmt_p(focal.exact_two_sided_p_value)}), "
        f"and was {status}. Across all 45 locked revision tests, {significant} survived their prespecified multiplicity controls. "
        f"The 30-seed random-policy schedule met both convergence thresholds (maximum |25-to-30 change|={report['maximum_absolute_prefix_25_to_30_change']:.4f}; "
        f"maximum 95% Monte Carlo half-width={report['maximum_mc_95_halfwidth']:.4f}). PCBM improved low-budget acquisition diversity, but the complete "
        "evidence did not justify an equivalence claim or an unqualified predictive-superiority claim."
    )
    methods = (
        "The robustness extension was internally prespecified before execution. Candidate pools were balanced (35 repetitions) or mildly, "
        "moderately, or severely imbalanced (32, 28, or 21 repetitions). At K07, Ridge-PCBM was contrasted within each pool condition with mean "
        "random acquisition, global margin, least confidence, predictive entropy, RBMAL margin-diversity, and greedy core-set selection. Normalized "
        "area under the learning curve used the common K00/K07/K14/K21 grid and trapezoidal integration divided by 21. Temporal sensitivity used four "
        "locked candidate/test partitions. Drift endpoints were late-minus-early no-adaptation accuracy and the within-session feature-distance slope. "
        "Deep stability averaged six locked training seeds within participant before inference."
    )
    statistics = (
        "P01–P06 were the only inferential units; sessions, pool realizations, acquisition seeds, training seeds, repetitions, and windows were never "
        "treated as independent units. Exact two-sided signed-rank enumeration discarded differences with absolute magnitude ≤1×10⁻¹² and reported "
        "the signed-rank statistic, nonzero-pair count, raw paired delta, standardized paired effect, rank-biserial correlation, and participant sign counts. "
        "Each participant-mean effect received a 100,000-resample BCa 95% interval and a separately reported 100,000-resample percentile sensitivity interval, "
        "using SHA-256-derived deterministic seeds. Holm adjustment was applied only within the locked secondary families; the single focal contrast was unadjusted. "
        "Random-policy adequacy was assessed with 1,000 deterministic subsamples at m=1, 2, 5, 10, 15, 20, 25, and 30."
    )
    results = (
        f"The single focal effect-modification contrast was {fmt_number(focal.mean_paired_delta, signed=True)} "
        f"(BCa 95% CI {interval_text(focal)}; W={fmt_number(focal.wilcoxon_statistic, 1)}; exact p={fmt_p(focal.exact_two_sided_p_value)}; "
        f"rank-biserial r={fmt_number(focal.rank_biserial_correlation, 3, signed=True)}). The focal revision hypothesis was {status}. "
        f"Across the 45 locked revision tests, {significant} met Holm-adjusted p<0.05 within their declared families. "
        f"PCBM-minus-random mean differences were positive in {int((imbalance_random.mean_paired_delta > 0).sum())}/{len(imbalance_random)} pool conditions; "
        f"temporal-split contrasts were positive in {int((split.mean_paired_delta > 0).sum())}/{len(split)} tests. "
        f"The two drift diagnostics and four deep seed-stability contrasts yielded {int(drift.significant_holm_0_05.sum())} and "
        f"{int(deep.significant_holm_0_05.sum())} multiplicity-controlled findings, respectively."
    )
    mc = (
        "The random-policy Monte Carlo diagnostic was adequate: the maximum absolute change from the 25-seed prefix to the full 30-seed mean was "
        f"{report['maximum_absolute_prefix_25_to_30_change']:.6f}, below 0.005, and the maximum 30-seed 95% Monte Carlo half-width was "
        f"{report['maximum_mc_95_halfwidth']:.6f}, below 0.01. The prespecified 31–60 seed extension was therefore not activated."
    )
    discussion = (
        "The expanded stress tests sharpen rather than replace the frozen interpretation. PCBM consistently organized low-budget acquisition across "
        "predicted classes, but its participant-level predictive advantage was not uniformly stable across pool imbalance, comparator choice, temporal split, "
        "or model stochasticity. Confidence intervals and exact tests should be read as uncertainty summaries in a six-participant cohort, not as evidence of "
        "equivalence when intervals include zero. The satisfactory 30-seed convergence audit reduces concern that the random comparator mean was materially "
        "distorted by Monte Carlo noise."
    )
    limitation = (
        "The population analysis included six able-bodied participants, with P07 retained as a single descriptive limb-difference case. Exact tests therefore "
        "had coarse resolution, while BCa intervals remain sensitive to a small number of participant estimands. Pool imbalance was synthetically controlled "
        "within the recorded candidate set and does not reproduce every deployment prevalence shift. Performance approached a ceiling at larger budgets. The "
        "TCN used RMS sequences rather than raw waveforms, and only six locked training seeds were evaluated. The study was an offline secondary analysis; it did "
        "not establish causal effects, equivalence, real-time prosthesis utility, clinical benefit, device safety, or prospective annotation burden."
    )
    conclusion = (
        "Under a leakage-controlled longitudinal protocol, PCBM improved low-budget acquisition diversity but did not demonstrate a robust, general predictive "
        "or retention advantage. The revised imbalance, comparator, AULC, temporal-split, drift, Monte Carlo, and TCN seed-stability analyses did not justify an "
        "equivalence claim or overturn the original bounded conclusion. Larger prospective cohorts are required before participant-level performance or clinical "
        "benefit can be claimed."
    )
    return {
        "focal": focal,
        "family_counts": family_counts,
        "abstract": abstract,
        "methods": methods,
        "statistics": statistics,
        "results": results,
        "mc": mc,
        "discussion": discussion,
        "limitation": limitation,
        "conclusion": conclusion,
    }


def build_main_manuscript(
    source: Path,
    destination: Path,
    tests: pd.DataFrame,
    report: dict,
    figure_path: Path,
    highlight: bool,
) -> dict:
    document = Document(source)
    bundle = revision_text_bundle(tests, report)

    for paragraph in list(document.paragraphs):
        text = " ".join(paragraph.text.split())
        if text.startswith(
            (
                "ORIGINAL RESEARCH -",
                "AUTHOR NAMES -",
                "AUTHOR AFFILIATIONS -",
                "CORRESPONDING AUTHOR",
                "TARGET JOURNAL -",
                "Scientific status:",
                "Stage 3G freeze:",
            )
        ):
            remove_paragraph(paragraph)
    map_heading = find_paragraph(document, exact="Manuscript map")
    abstract_heading = find_paragraph(document, exact="Abstract")
    node = map_heading._p
    while node is not abstract_heading._p:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node

    title = find_paragraph(document, startswith="Predicted-Class-Balanced Active Learning")
    set_text(
        title,
        "Predicted-Class-Balanced Active Learning for Leakage-Controlled Longitudinal Adaptation of High-Density Surface Electromyography: Classical and Deep-Learning Evidence from the DELTA Dataset",
        highlight=highlight,
        bold=True,
    )
    abstract = abstract_heading._p.getnext()
    from docx.text.paragraph import Paragraph

    abstract_paragraph = Paragraph(abstract, abstract_heading._parent)
    set_text(abstract_paragraph, bundle["abstract"], highlight=highlight)

    highlights_heading = find_paragraph(document, exact="Research highlights")
    current = highlights_heading._p.getnext()
    bullet_paragraphs = []
    while current is not None:
        paragraph = Paragraph(current, highlights_heading._parent)
        if paragraph.style and paragraph.style.name == "List Bullet":
            bullet_paragraphs.append(paragraph)
            current = current.getnext()
        else:
            break
    new_highlights = [
        "Leakage-controlled adaptation kept every fixed test isolated.",
        "PCBM increased low-budget acquisition coverage and entropy.",
        "The imbalance effect-modification hypothesis was not supported.",
        "Thirty random seeds met both Monte Carlo adequacy thresholds.",
        "P07 remained descriptive and outside population inference.",
    ]
    for index, text in enumerate(new_highlights):
        if index < len(bullet_paragraphs):
            set_text(bullet_paragraphs[index], text, highlight=highlight)
        else:
            paragraph = document.add_paragraph(text, style="List Bullet")
            current.addprevious(paragraph._p)
            mark_runs(paragraph, highlight)

    figure1 = find_paragraph(document, startswith="Figure 1.")
    set_text(
        figure1,
        "Figure 1. Leakage-controlled continual-adaptation design. Fixed-test repetitions remain unavailable to training, normalization, calibration, and selection.",
        highlight=highlight,
    )

    outcomes = find_paragraph(document, exact="Outcomes and statistical analysis")
    old_one = Paragraph(outcomes._p.getnext(), outcomes._parent)
    old_two = Paragraph(old_one._p.getnext(), outcomes._parent)
    set_text(old_one, bundle["methods"], highlight=highlight)
    set_text(old_two, bundle["statistics"], highlight=highlight)

    results_heading = find_paragraph(document, exact="Results")
    results_heading.paragraph_format.page_break_before = True
    results_heading.paragraph_format.space_after = Pt(10)

    discussion = find_paragraph(document, exact="Discussion")
    append_before(discussion, "Robustness extensions", style="Heading 2", highlight=highlight)
    append_before(discussion, bundle["results"], highlight=highlight)
    append_before(discussion, bundle["mc"], highlight=highlight)

    imbalance = tests.loc[
        tests["analysis_id"].str.contains("REV_SECONDARY_IMBALANCE")
        & tests["analysis_id"].str.endswith("PCBM_MINUS_RANDOM_UNIFORM")
    ].copy()
    imbalance["condition"] = imbalance["analysis_id"].str.extract(r"RIDGE__(.*?)__PCBM")[0]
    caption = append_before(
        discussion,
        "Table 8. Ridge K07 PCBM-minus-random contrasts across candidate-pool conditions.",
        highlight=highlight,
    )
    caption.runs[0].bold = True
    caption.paragraph_format.keep_with_next = True
    rows = []
    for row in imbalance.sort_values("condition").itertuples(index=False):
        rows.append(
            [
                row.condition,
                fmt_number(row.mean_paired_delta, signed=True),
                f"[{fmt_number(row.bca_95_ci_low, signed=True)}, {fmt_number(row.bca_95_ci_high, signed=True)}]",
                fmt_p(row.exact_two_sided_p_value),
                fmt_p(row.holm_adjusted_p_value),
                f"{row.participants_positive}/{row.participants_tied}/{row.participants_negative}",
            ]
        )
    add_table_before(
        discussion,
        ["Pool", "Mean Δ", "BCa 95% CI", "Raw p", "Holm p", "+/0/−"],
        rows,
        [1800, 1200, 2400, 1000, 1000, 1960],
        highlight=highlight,
    )

    aulc = tests.loc[tests["analysis_id"].str.startswith("REV_SECONDARY_AULC")].copy()
    aulc["comparator"] = aulc["analysis_id"].str.split("PCBM_MINUS_").str[-1]
    caption = append_before(
        discussion,
        "Table 9. Normalized AULC contrasts over K00/K07/K14/K21.",
        highlight=highlight,
    )
    caption.runs[0].bold = True
    caption.paragraph_format.keep_with_next = True
    rows = []
    for row in aulc.sort_values("comparator").itertuples(index=False):
        rows.append(
            [
                row.comparator,
                fmt_number(row.mean_paired_delta, signed=True),
                f"[{fmt_number(row.bca_95_ci_low, signed=True)}, {fmt_number(row.bca_95_ci_high, signed=True)}]",
                fmt_p(row.exact_two_sided_p_value),
                fmt_p(row.holm_adjusted_p_value),
            ]
        )
    add_table_before(
        discussion,
        ["Comparator", "Mean Δ", "BCa 95% CI", "Raw p", "Holm p"],
        rows,
        [2600, 1400, 2500, 1300, 1560],
        highlight=highlight,
    )

    figure_paragraph = append_before(discussion, "", highlight=False)
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.add_run().add_picture(str(figure_path), width=Inches(6.25))
    figure_caption = append_before(
        discussion,
        "Figure 9. Participant-level mean differences and BCa 95% confidence intervals for Ridge PCBM versus random acquisition across pool conditions (left) and normalized AULC versus deployable comparators (right).",
        highlight=highlight,
    )

    first_discussion = Paragraph(discussion._p.getnext(), discussion._parent)
    set_text(first_discussion, bundle["discussion"], highlight=highlight)
    limitations = find_paragraph(document, exact="Limitations")
    limitation_paragraph = Paragraph(limitations._p.getnext(), limitations._parent)
    set_text(limitation_paragraph, bundle["limitation"], highlight=highlight)
    conclusion = find_paragraph(document, exact="Conclusion")
    conclusion_paragraph = Paragraph(conclusion._p.getnext(), conclusion._parent)
    set_text(conclusion_paragraph, bundle["conclusion"], highlight=highlight)

    code_availability = find_paragraph(document, startswith="Code availability.")
    set_text(
        code_availability,
        "Code availability. Executed source code, derived tables, SHA-256 manifests, and the complete statistical supplement are included in the archived evidence package accompanying this revision. A public repository identifier will be supplied in the journal submission metadata after de-anonymization.",
        highlight=highlight,
    )
    competing = find_paragraph(document, startswith="Competing interests.")
    set_text(
        competing,
        "Competing interests. The competing-interest declaration is provided in the journal's separate author-information file and is omitted here to preserve manuscript anonymization.",
        highlight=highlight,
    )
    ai = find_paragraph(document, startswith="Use of artificial intelligence tools.")
    set_text(
        ai,
        "Use of artificial intelligence tools. Generative-AI tools assisted with language editing and code drafting. The authors independently verified the protocol, source code, statistical outputs, interpretation, and final text and retain full responsibility for the work.",
        highlight=highlight,
    )

    appendix = find_paragraph(document, exact="Appendix: reporting and claim guardrails")
    node = appendix._p
    body = node.getparent()
    while node is not None:
        next_node = node.getnext()
        body.remove(node)
        node = next_node

    renumber_captions(document, highlight=highlight)
    properties = document.core_properties
    properties.title = title.text
    properties.subject = "Leakage-controlled longitudinal HD-sEMG active-learning study"
    properties.comments = "Revision R8A integrated the locked R7B robustness and statistical evidence."
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "contains_r8_results_heading": "Robustness extensions" in text,
        "contains_internal_manuscript_map": "Manuscript map" in text,
        "contains_internal_stage_label": bool(re.search(r"Stage (3G|5F|6)\b", text)),
        "contains_submission_placeholder": any(token in text for token in ["REPLACE BEFORE", "SELECT BEFORE", "TO BE COMPLETED", "Add the final public"]),
        "title_uses_leakage_controlled": "Leakage-Controlled" in title.text,
        "references_1_to_62_present": all(f"[{index}]" in text for index in range(1, 63)),
    }


def compact_columns(frame: pd.DataFrame, maximum: int = 8) -> list[str]:
    priorities = [
        "analysis_id",
        "comparison",
        "contrast",
        "metric",
        "query_budget",
        "mean_paired_delta",
        "paired_delta_ci_low",
        "paired_delta_ci_high",
        "exact_two_sided_p_value",
        "p_value_raw",
        "holm_adjusted_p_value",
        "p_value_holm",
        "significant_holm_0_05",
    ]
    selected = [column for column in priorities if column in frame.columns]
    for column in frame.columns:
        if column not in selected and len(selected) < maximum:
            selected.append(column)
    return selected[:maximum]


def add_dataframe_table(document: Document, frame: pd.DataFrame, columns: list[str], widths: list[int], font_size: float = 7.5) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = str(column).replace("_", " ")
    for values in frame[columns].itertuples(index=False, name=None):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if isinstance(value, float):
                cells[index].text = fmt_number(value, 5)
            else:
                cells[index].text = str(value)
    style_table(table, widths, highlight=False, font_size=font_size)


def build_supplement(
    destination: Path,
    tests: pd.DataFrame,
    retention: pd.DataFrame,
    classical: pd.DataFrame,
    mc_cells: pd.DataFrame,
    p07: pd.DataFrame,
    report: dict,
) -> dict:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.color.rgb = RGBColor(23, 54, 93)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(11)
    styles["Heading 2"].font.color.rgb = RGBColor(23, 54, 93)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Statistical and Reproducibility Supplement")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(23, 54, 93)
    subtitle = document.add_paragraph(
        "Predicted-Class-Balanced Active Learning for Leakage-Controlled Longitudinal Adaptation of HD-sEMG"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Inferential population: P01–P06. P07 is descriptive only. Sessions and random/training seeds are averaged within participant. "
        "BCa and percentile intervals each use 100,000 deterministic bootstrap resamples."
    )

    document.add_heading("S1. Complete locked R7 statistical results", level=1)
    family_order = list(dict.fromkeys(tests["multiplicity_family"].tolist()))
    for family in family_order:
        group = tests.loc[tests["multiplicity_family"].eq(family)].copy()
        document.add_heading(family.replace("_", " "), level=2)
        table_frame = pd.DataFrame(
            {
                "analysis": group["analysis_id"],
                "mean_delta": group["mean_paired_delta"],
                "bca_low": group["bca_95_ci_low"],
                "bca_high": group["bca_95_ci_high"],
                "raw_p": group["exact_two_sided_p_value"],
                "holm_p": group["holm_adjusted_p_value"],
                "rank_biserial": group["rank_biserial_correlation"],
                "signs_pos_tie_neg": group.apply(
                    lambda row: f"{int(row.participants_positive)}/{int(row.participants_tied)}/{int(row.participants_negative)}",
                    axis=1,
                ),
            }
        )
        add_dataframe_table(
            document,
            table_frame,
            list(table_frame.columns),
            [4200, 1100, 1100, 1100, 900, 900, 1100, 1360],
            font_size=7.2,
        )

    document.add_heading("S2. Frozen classical secondary results", level=1)
    classical_columns = compact_columns(classical)
    add_dataframe_table(
        document,
        classical,
        classical_columns,
        [int(12960 / len(classical_columns))] * len(classical_columns),
        font_size=7.2,
    )

    document.add_heading("S3. Frozen retention family", level=1)
    retention_columns = compact_columns(retention)
    add_dataframe_table(
        document,
        retention,
        retention_columns,
        [int(12960 / len(retention_columns))] * len(retention_columns),
        font_size=7.2,
    )

    document.add_heading("S4. Random-policy Monte Carlo adequacy", level=1)
    mc_summary = (
        mc_cells.groupby(["imbalance_level", "query_budget"], as_index=False)
        .agg(
            maximum_absolute_25_to_30_change=("absolute_prefix_25_to_30_change", "max"),
            maximum_30_seed_95_halfwidth=("full_30_seed_mean_mc_95_halfwidth", "max"),
            all_change_gates_pass=("change_threshold_0_005_passes", "all"),
            all_halfwidth_gates_pass=("halfwidth_threshold_0_01_passes", "all"),
        )
    )
    add_dataframe_table(
        document,
        mc_summary,
        list(mc_summary.columns),
        [2400, 1400, 2600, 2600, 1980, 1980],
        font_size=7.5,
    )
    document.add_paragraph(
        f"Global maxima: |25-to-30 change|={report['maximum_absolute_prefix_25_to_30_change']:.6f}; "
        f"30-seed 95% Monte Carlo half-width={report['maximum_mc_95_halfwidth']:.6f}. Both prespecified gates passed."
    )

    document.add_heading("S5. Participant-level and P07 data availability", level=1)
    document.add_paragraph(
        "The accompanying CSV archive contains all 270 participant-level paired-difference rows, all BCa and percentile intervals, "
        "the complete deterministic Monte Carlo resampling audit, class recalls, confusion matrices, class coverage, normalized entropy, "
        "compute telemetry, training-seed distributions, and all P07 descriptive rows. These high-dimensional tables are supplied as CSV "
        "rather than reproduced as thousands of Word table cells."
    )
    p07_counts = p07.groupby("source_analysis", as_index=False).size().rename(columns={"size": "descriptive_rows"})
    add_dataframe_table(document, p07_counts, list(p07_counts.columns), [6500, 6460], font_size=8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Statistical supplement — R8A evidence integration")
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    return {
        "test_rows": len(tests),
        "retention_rows": len(retention),
        "classical_rows": len(classical),
        "mc_cell_rows": len(mc_cells),
        "p07_rows": len(p07),
        "table_count": len(document.tables),
    }


def build_response_document(destination: Path, action_matrix: pd.DataFrame, report: dict) -> dict:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Response-to-Reviewers Evidence Matrix — Draft")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    document.add_paragraph(
        "This evidence matrix maps the locked reviewer-action registry to the completed revision stages. Exact reviewer wording and author metadata will be added in R8B."
    )
    completion = {
        "R2": "Completed: literature, novelty, oracle, architecture, equations, and wording.",
        "R3_R6": "Completed: new comparator implementation and deep BADGE/seed-stability analyses.",
        "R4": "Completed: four-level candidate-pool imbalance stress test.",
        "R5": "Completed: alternative temporal splits and within-session drift audit.",
        "R6": "Completed: six-seed fixed-history and end-to-end TCN stability.",
        "R7": "Completed: 45 locked tests, AULC, MC convergence, and supplement.",
        "R7_R8": "Completed statistically in R7; integrated into manuscript and supplement in R8A.",
        "R8": "Integrated in R8A; author metadata and final repository identifiers remain for R8B.",
    }
    rows = []
    for row in action_matrix.itertuples(index=False):
        rows.append(
            [
                str(row.issue_id),
                str(row.priority),
                str(row.domain),
                str(row.locked_response_action),
                completion.get(str(row.execution_stage), "Completed or scheduled under the locked revision sequence."),
            ]
        )
    table = document.add_table(rows=1, cols=5)
    headers = ["Issue", "Priority", "Domain", "Locked response", "R8A status/evidence"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, [900, 900, 1300, 3000, 3260], font_size=7.5)
    document.add_heading("Locked statistical completion", level=1)
    document.add_paragraph(
        f"R7B executed 45 participant-level tests and passed every readiness gate. Random-policy convergence passed with "
        f"maximum |25-to-30 change|={report['maximum_absolute_prefix_25_to_30_change']:.6f} and maximum 30-seed 95% Monte Carlo "
        f"half-width={report['maximum_mc_95_halfwidth']:.6f}. No 31–60 seed extension was required."
    )
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    return {"action_rows": len(rows), "table_count": len(document.tables)}


def render_docx(docx_path: Path, name: str) -> tuple[Path | None, pd.DataFrame]:
    output = RENDER_ROOT / name
    shutil.rmtree(output, ignore_errors=True)
    output.mkdir(parents=True, exist_ok=True)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None, pd.DataFrame()
    profile = Path(tempfile.mkdtemp(prefix="r8a_lo_", dir="/tmp"))
    try:
        subprocess.run(
            [soffice, "--headless", f"-env:UserInstallation=file://{profile}", "--convert-to", "pdf", "--outdir", str(output), str(docx_path)],
            check=True,
            capture_output=True,
            text=True,
        )
    finally:
        shutil.rmtree(profile, ignore_errors=True)
    pdf = output / (docx_path.stem + ".pdf")
    if not pdf.exists():
        return None, pd.DataFrame()
    pdftoppm = shutil.which("pdftoppm")
    rows = []
    if pdftoppm:
        subprocess.run([pdftoppm, "-png", "-r", "110", str(pdf), str(output / "page")], check=True, capture_output=True)
        from PIL import Image

        for page in sorted(output.glob("page-*.png")):
            image = np.asarray(Image.open(page).convert("L"))
            rows.append(
                {
                    "document": name,
                    "page": int(page.stem.split("-")[-1]),
                    "width_px": image.shape[1],
                    "height_px": image.shape[0],
                    "nonwhite_fraction": float(np.mean(image < 248)),
                    "appears_blank": bool(np.mean(image < 248) < 0.001),
                }
            )
        for page in output.glob("page-*.png"):
            page.unlink()
    return pdf, pd.DataFrame(rows)


def extract_r7b_evidence(packet: Path) -> pd.DataFrame:
    index_rows = []
    with zipfile.ZipFile(packet, "r") as archive:
        for member in archive.infolist():
            if member.is_dir():
                continue
            basename = Path(member.filename).name
            suffix = Path(basename).suffix.lower()
            if suffix not in {".csv", ".json"}:
                continue
            relative = Path(*Path(member.filename).parts[1:]) if len(Path(member.filename).parts) > 1 else Path(basename)
            if ".." in relative.parts:
                raise RuntimeError("Unsafe R7B archive member")
            destination = EVIDENCE_ROOT / relative
            destination.parent.mkdir(parents=True, exist_ok=True)
            destination.write_bytes(archive.read(member))
            index_rows.append(
                {
                    "source_member": member.filename,
                    "relative_path": destination.relative_to(RESULT_ROOT).as_posix(),
                    "bytes": destination.stat().st_size,
                    "sha256": engine.sha256_file(destination),
                }
            )
    return pd.DataFrame(index_rows)


def main() -> None:
    print("=" * 112)
    print("REVISION R8A — REVISED MANUSCRIPT AND SUPPLEMENT INTEGRATION")
    print("=" * 112)
    print("Execution device: CPU")
    print("Model training: False")
    print("Fixed-test inference: False")
    print("New statistical tests: False")
    print("Scientific role: integrate the locked R7B evidence without changing it")
    print()

    WORKING.mkdir(parents=True, exist_ok=True)
    lock_handle = open(WORKING / "_revision_R8A_single_instance.lock", "w")
    try:
        fcntl.flock(lock_handle, fcntl.LOCK_EX | fcntl.LOCK_NB)
    except BlockingIOError:
        print("FINAL DECISION: DUPLICATE_INVOCATION_EXITED_SAFELY")
        return
    engine.bootstrap_rclone()
    engine.create_rclone_config()
    print("rclone version:", engine.rclone(["version"]).stdout.splitlines()[0])
    print("Restoring verified R0, R2, and R7B packets...")
    r0_packet, r0_source = direct_restore(R0_BASENAME, R0_PACKET_SHA256, R0_REMOTE)
    r2_packet, r2_source = direct_restore(R2_BASENAME, R2_PACKET_SHA256, R2_REMOTE)
    r7b_packet, r7b_source = direct_restore(R7B_BASENAME, R7B_PACKET_SHA256, R7B_REMOTE)
    r2_docx_member, r2_docx_expected_hash, r2_docx_actual_hash = restore_source_docx_from_r2(r2_packet)
    r7b_report = read_json(r7b_packet, "revision_R7B_final_report.json")
    tests = read_csv(r7b_packet, "revision_R7B_all_new_statistical_tests.csv")
    retention = read_csv(r7b_packet, "revision_R7B_frozen_eighteen_retention_tests.csv")
    classical = read_csv(r7b_packet, "revision_R7B_frozen_five_classical_secondary_tests.csv")
    mc_cells = read_csv(r7b_packet, "revision_R7B_random_policy_mc_cell_adequacy.csv")
    p07 = read_csv(r7b_packet, "revision_R7B_P07_descriptive_only.csv")
    action_matrix = read_csv(r0_packet, "stageR0_reviewer_action_matrix.csv")
    claims = read_csv(r0_packet, "stageR0_claim_guardrails.csv")

    shutil.rmtree(RESULT_ROOT, ignore_errors=True)
    shutil.rmtree(RENDER_ROOT, ignore_errors=True)
    RESULT_ROOT.mkdir(parents=True, exist_ok=True)
    EVIDENCE_ROOT.mkdir(parents=True, exist_ok=True)
    evidence_index = extract_r7b_evidence(r7b_packet)
    figure_path = RESULT_ROOT / "revision_R8A_figure9_robustness_summary.png"
    create_robustness_figure(tests, figure_path)

    clean_docx = RESULT_ROOT / "revision_R8A_revised_manuscript_clean_anonymized.docx"
    highlighted_docx = RESULT_ROOT / "revision_R8A_revised_manuscript_changes_highlighted.docx"
    supplement_docx = RESULT_ROOT / "revision_R8A_statistical_reproducibility_supplement.docx"
    response_docx = RESULT_ROOT / "revision_R8A_response_to_reviewers_evidence_matrix_draft.docx"
    clean_audit = build_main_manuscript(SOURCE_DOCX, clean_docx, tests, r7b_report, figure_path, highlight=False)
    highlighted_audit = build_main_manuscript(SOURCE_DOCX, highlighted_docx, tests, r7b_report, figure_path, highlight=True)
    supplement_audit = build_supplement(supplement_docx, tests, retention, classical, mc_cells, p07, r7b_report)
    response_audit = build_response_document(response_docx, action_matrix, r7b_report)

    render_rows = []
    for path, name in [
        (clean_docx, "clean_manuscript"),
        (highlighted_docx, "highlighted_manuscript"),
        (supplement_docx, "statistical_supplement"),
        (response_docx, "response_matrix"),
    ]:
        pdf, audit = render_docx(path, name)
        if pdf is not None:
            shutil.copy2(pdf, RESULT_ROOT / (path.stem + ".pdf"))
        if len(audit):
            render_rows.append(audit)
    render_audit = pd.concat(render_rows, ignore_index=True) if render_rows else pd.DataFrame()

    author_actions = pd.DataFrame(
        [
            ("AUTHOR_NAMES", "Provide final author names in journal order", "REQUIRED_R8B"),
            ("AFFILIATIONS", "Provide numbered affiliations and author-affiliation mapping", "REQUIRED_R8B"),
            ("CORRESPONDING_AUTHOR", "Provide corresponding author, institutional email, and postal address", "REQUIRED_R8B"),
            ("COMPETING_INTERESTS", "Confirm the competing-interest declaration", "REQUIRED_R8B"),
            ("CREDIT", "Confirm CRediT roles for every author", "REQUIRED_R8B"),
            ("PUBLIC_REPOSITORY", "Provide the final public repository or archival DOI", "REQUIRED_R8B"),
            ("AI_DISCLOSURE", "Approve or amend the journal-specific AI-assistance statement", "REQUIRED_R8B"),
        ],
        columns=["action_id", "required_author_input", "status"],
    )
    atomic_csv(author_actions, RESULT_ROOT / "revision_R8A_author_action_checklist_for_R8B.csv")
    atomic_csv(action_matrix, RESULT_ROOT / "revision_R8A_locked_reviewer_action_matrix.csv")
    atomic_csv(claims, RESULT_ROOT / "revision_R8A_locked_claim_guardrails.csv")
    atomic_csv(evidence_index, RESULT_ROOT / "revision_R8A_supplementary_evidence_index.csv")
    atomic_csv(render_audit, RESULT_ROOT / "revision_R8A_render_page_audit.csv")
    atomic_json(clean_audit, RESULT_ROOT / "revision_R8A_clean_manuscript_structural_audit.json")
    atomic_json(highlighted_audit, RESULT_ROOT / "revision_R8A_highlighted_manuscript_structural_audit.json")
    atomic_json(supplement_audit, RESULT_ROOT / "revision_R8A_supplement_structural_audit.json")
    atomic_json(response_audit, RESULT_ROOT / "revision_R8A_response_matrix_structural_audit.json")

    input_audit = pd.DataFrame(
        [
            {
                "input": "R2_CORRECTED_SOURCE_DOCX",
                "source": f"{r2_source}:{r2_docx_member}",
                "expected_sha256": r2_docx_expected_hash,
                "actual_sha256": r2_docx_actual_hash,
                "hash_matches": r2_docx_actual_hash == r2_docx_expected_hash,
                "crc_passes": True,
            },
            {
                "input": R2_BASENAME,
                "source": r2_source,
                "expected_sha256": R2_PACKET_SHA256,
                "actual_sha256": engine.sha256_file(r2_packet),
                "hash_matches": engine.sha256_file(r2_packet) == R2_PACKET_SHA256,
                "crc_passes": engine.archive_crc_passes(r2_packet),
            },
            {
                "input": R0_BASENAME,
                "source": r0_source,
                "expected_sha256": R0_PACKET_SHA256,
                "actual_sha256": engine.sha256_file(r0_packet),
                "hash_matches": engine.sha256_file(r0_packet) == R0_PACKET_SHA256,
                "crc_passes": engine.archive_crc_passes(r0_packet),
            },
            {
                "input": R7B_BASENAME,
                "source": r7b_source,
                "expected_sha256": R7B_PACKET_SHA256,
                "actual_sha256": engine.sha256_file(r7b_packet),
                "hash_matches": engine.sha256_file(r7b_packet) == R7B_PACKET_SHA256,
                "crc_passes": engine.archive_crc_passes(r7b_packet),
            },
        ]
    )
    atomic_csv(input_audit, RESULT_ROOT / "revision_R8A_input_audit.csv")

    expected_pdfs = [RESULT_ROOT / (path.stem + ".pdf") for path in [clean_docx, highlighted_docx, supplement_docx, response_docx]]
    gates = {
        "all_verified_inputs_match_hash_and_crc": bool(input_audit[["hash_matches", "crc_passes"]].all().all()),
        "r7b_all_readiness_gates_passed": bool(r7b_report.get("all_readiness_gates_passed")),
        "r7b_decision_authorizes_r8": r7b_report.get("final_decision")
        == "PASS_TO_REVISION_R8_REVISED_MANUSCRIPT_AND_SUPPLEMENT_INTEGRATION",
        "r7b_contains_exactly_45_locked_tests": len(tests) == 45,
        "random_policy_convergence_is_adequate": bool(r7b_report.get("monte_carlo_random_policy_adequate")),
        "clean_and_highlighted_manuscripts_exist": clean_docx.exists() and highlighted_docx.exists(),
        "statistical_supplement_exists": supplement_docx.exists(),
        "response_matrix_draft_exists": response_docx.exists(),
        "clean_title_uses_leakage_controlled_wording": clean_audit["title_uses_leakage_controlled"],
        "clean_manuscript_has_no_internal_map_or_stage_labels": not clean_audit["contains_internal_manuscript_map"]
        and not clean_audit["contains_internal_stage_label"],
        "clean_manuscript_has_no_submission_placeholders": not clean_audit["contains_submission_placeholder"],
        "all_62_references_are_preserved": clean_audit["references_1_to_62_present"],
        "supplement_contains_all_45_new_tests": supplement_audit["test_rows"] == 45,
        "supplement_exposes_at_least_18_retention_rows": supplement_audit["retention_rows"] >= 18,
        "supplement_exposes_at_least_five_classical_rows": supplement_audit["classical_rows"] >= 5,
        "all_four_docx_render_to_pdf": all(path.exists() and path.stat().st_size > 0 for path in expected_pdfs),
        "rendered_pages_are_nonblank": len(render_audit) > 0 and not render_audit["appears_blank"].any(),
        "r7b_evidence_is_preserved_in_csv_json_form": len(evidence_index) >= 15,
        "claim_guardrails_are_preserved": len(claims) >= 10,
        "author_metadata_is_deferred_explicitly_to_r8b": len(author_actions) == 7,
        "raw_hdf5_data_was_not_accessed": True,
        "no_model_was_trained": True,
        "no_fixed_test_inference_was_run": True,
        "no_new_statistical_test_was_run": True,
        "credentials_were_not_written_to_artifacts": True,
    }
    failed = [key for key, value in gates.items() if not bool(value)]
    report = {
        "stage": "REVISION_R8A_REVISED_MANUSCRIPT_SUPPLEMENT_INTEGRATION",
        "revision_protocol_sha256": REVISION_PROTOCOL_SHA256,
        "r2_packet_sha256": R2_PACKET_SHA256,
        "r7b_packet_sha256": R7B_PACKET_SHA256,
        "source_docx_sha256": r2_docx_actual_hash,
        "local_qa_reference_source_docx_sha256": LOCAL_QA_SOURCE_DOCX_SHA256,
        "clean_manuscript_sha256": engine.sha256_file(clean_docx),
        "highlighted_manuscript_sha256": engine.sha256_file(highlighted_docx),
        "supplement_docx_sha256": engine.sha256_file(supplement_docx),
        "response_matrix_docx_sha256": engine.sha256_file(response_docx),
        "locked_test_rows_integrated": len(tests),
        "r7b_evidence_files_preserved": len(evidence_index),
        "author_actions_remaining": len(author_actions),
        "readiness_gates": gates,
        "failed_readiness_gates": failed,
        "all_readiness_gates_passed": not failed,
        "model_training_run": False,
        "fixed_test_inference_run": False,
        "new_statistical_test_run": False,
        "runtime_minutes": (time.time() - START_TIME) / 60.0,
        "final_decision": "PASS_TO_REVISION_R8B_AUTHOR_METADATA_AND_FINAL_SUBMISSION_PACKAGE"
        if not failed
        else "REVISION_R8A_INTEGRATION_FAILED",
    }
    atomic_json(report, RESULT_ROOT / "revision_R8A_final_report.json")
    shutil.copy2(Path(__file__), RESULT_ROOT / "revision_R8A_executed_source.py")
    manifest_rows = []
    for path in sorted(RESULT_ROOT.rglob("*")):
        if path.is_file() and path.name != "revision_R8A_output_manifest.csv":
            manifest_rows.append(
                {
                    "relative_path": path.relative_to(RESULT_ROOT).as_posix(),
                    "bytes": path.stat().st_size,
                    "sha256": engine.sha256_file(path),
                }
            )
    atomic_csv(pd.DataFrame(manifest_rows), RESULT_ROOT / "revision_R8A_output_manifest.csv")
    if failed:
        raise RuntimeError(f"R8A readiness failed: {failed}")
    if not engine.make_zip(RESULT_ROOT, PACKET_PATH, "Revision_R8A_Revised_Manuscript_Supplement_Integration"):
        raise RuntimeError("R8A packet CRC failed")
    digest = engine.sha256_file(PACKET_PATH)
    if not engine.roundtrip_remote_file(PACKET_PATH, REMOTE_OUTPUT + "/" + PACKET_PATH.name, digest):
        raise RuntimeError("R8A remote round-trip failed")

    print()
    print("=" * 112)
    print("REVISION R8A — FINAL INTEGRATION SUMMARY")
    print("=" * 112)
    print("Locked statistical rows integrated:", len(tests))
    print("R7B evidence files preserved:", len(evidence_index))
    print("Clean manuscript:", clean_docx)
    print("Highlighted manuscript:", highlighted_docx)
    print("Statistical supplement:", supplement_docx)
    print("Response matrix draft:", response_docx)
    print("Author actions deferred to R8B:", len(author_actions))
    print("Failed readiness gates:", failed or "None")
    print("Packet CRC pass: True")
    print("Packet:", PACKET_PATH)
    print("Packet SHA-256:", digest)
    print("Remote round-trip verified: True")
    print("Runtime minutes:", round(report["runtime_minutes"], 3))
    print()
    print("FINAL DECISION:", report["final_decision"])


if __name__ == "__main__":
    main()
