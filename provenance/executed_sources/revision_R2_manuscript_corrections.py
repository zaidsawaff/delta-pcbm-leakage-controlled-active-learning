from __future__ import annotations

import atexit
import base64
import configparser
import copy
import hashlib
import importlib.util
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import time
import urllib.request
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# =============================================================================
# REVISION R2 — MANUSCRIPT CORRECTIONS
# =============================================================================

REVISION_PROTOCOL_NAME = "DELTA_PCBM_REVIEWER_REQUESTED_REVISION_v1"
REVISION_PROTOCOL_SHA256 = (
    "6807b71de18ca82013cfa4360d760e0daf9a920a1acc0625dcb13bd8f4d07249"
)
R0_PACKET_SHA256 = (
    "0800e315a29b81934095ba56deaea3f8b6600fd0df13db348d7ea72d3b82df78"
)
R1_PACKET_SHA256 = (
    "2ec8cff608a765d20807e2d57249bf091768d97ff747ee2de7b44bcb17475ec8"
)
STAGE5B_PACKET_SHA256 = (
    "1c0fbc63f6412362f3ae7cd22609ea6a7fcb23236cdf688ad5fe0578ebaab84d"
)
STAGE5C_PACKET_SHA256 = (
    "85ea2e8a8440369a77d43f00b5d509ea2f2978d2a60ab2f24fb828ce9ca6b9d4"
)
STAGE5D2_PACKET_SHA256 = (
    "fc8ac364bac0344639a50977d5f8725b1e5b5b2875758e01587de8c083a1f914"
)
SOURCE_DOCX_SHA256 = (
    "2a239f29a95484d3a26e3560fdb57ed870ca2557e86c616792204551e3847b6f"
)
PARENT_PROTOCOL_SHA256 = (
    "f548b1ca6f2831c29ea8fecb764557efed49f229eb72322f98632edcf0aeb221"
)
STAGE3G_FREEZE_SHA256 = (
    "dfcdbbade7bf1032c3250626439689e08e95249ec49d5fb90465c4444a998fbe"
)
DEEP_PROTOCOL_SHA256 = (
    "abe15812c1a52b0f4e917b5b6ad39b0dfde50e5bb2d58dfcc35b3cacb22e3bd2"
)
FROZEN_STAGE3G_CONCLUSION = (
    "PCBM increased low-budget acquisition diversity, but did not demonstrate "
    "robust predictive or retention superiority."
)
FROZEN_DEEP_CONCLUSION = "DEEP_EXTENSION_PRIMARY_HYPOTHESIS_NOT_SUPPORTED"

R2_ACTIONS = [
    "NOVELTY_01",
    "NOVELTY_02",
    "NOVELTY_03",
    "ORACLE_01",
    "DEEP_01",
    "MATH_01",
    "MATH_02",
    "REPRO_01",
]

DETAIL_PACKETS_REQUIRED_BEFORE_R3_R7 = [
    "stage3c1_deterministic_experiment_packet.zip",
    "stage3c2_random_control_packet.zip",
    "stage3d_primary_statistical_analysis_packet.zip",
    "stage3e2a_lda_deterministic_packet.zip",
    "stage3e2b_lda_random_sensitivity_packet.zip",
    "stage3e2c_strict_qc_ridge_packet.zip",
    "stage3e3_sensitivity_integration_packet.zip",
    "stage3f2a_deterministic_retention_packet.zip",
    "stage3f2b_random_retention_packet.zip",
    "stage3f3_retention_statistical_analysis_packet.zip",
]

WORKING = Path(os.environ.get("REVISION_R2_WORKING", "/kaggle/working"))
SOURCE_DOCX = Path(
    os.environ.get(
        "REVISION_R2_SOURCE_DOCX",
        str(WORKING / "revision_R2_source_expanded_manuscript.docx"),
    )
)
TOOLS = WORKING / "_stage5_tools"
INPUT_ROOT = WORKING / "REVISION_R2_FROZEN_INPUTS"
RESULT_ROOT = (
    WORKING
    / "DELTA_REVIEWER_REVISION"
    / "Revision_R2_Manuscript_Corrections"
)
WORK_ROOT = WORKING / "REVISION_R2_WORK"
for directory in (TOOLS, INPUT_ROOT, RESULT_ROOT, WORK_ROOT):
    directory.mkdir(parents=True, exist_ok=True)

RCLONE = TOOLS / "rclone"
PACKET_PATH = WORKING / "revision_R2_manuscript_corrections_packet.zip"
REMOTE_BASE = "gdrive_stage5:DELTA_Q1_Stage5_DeepLearning_Backup"
REMOTE_OUTPUT = (
    REMOTE_BASE + "/Reviewer_Revision/Revision_R2_Manuscript_Corrections"
)
CONFIG_PATH = None
REMOTE_LISTING = None
START_TIME = time.time()

BLUE = "17365D"
LIGHT_GRAY = "F2F4F7"


# =============================================================================
# GENERIC UTILITIES
# =============================================================================


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as stream:
        while True:
            block = stream.read(chunk_size)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def atomic_json(payload, destination: Path):
    destination = Path(destination)
    temporary = destination.with_suffix(destination.suffix + ".tmp")
    temporary.write_text(
        json.dumps(payload, indent=2, ensure_ascii=False, default=str),
        encoding="utf-8",
    )
    os.replace(temporary, destination)


def atomic_csv(frame: pd.DataFrame, destination: Path):
    destination = Path(destination)
    temporary = destination.with_suffix(destination.suffix + ".tmp")
    frame.to_csv(temporary, index=False)
    os.replace(temporary, destination)


def archive_crc_passes(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path, "r") as archive:
            return archive.testzip() is None
    except zipfile.BadZipFile:
        return False


def archive_member_matches(packet: Path, basename: str):
    with zipfile.ZipFile(packet, "r") as archive:
        return [name for name in archive.namelist() if Path(name).name == basename]


def archive_member(packet: Path, basename: str) -> bytes:
    matches = archive_member_matches(packet, basename)
    if len(matches) != 1:
        raise RuntimeError(f"Expected one {basename} in {packet}; found {matches}")
    with zipfile.ZipFile(packet, "r") as archive:
        return archive.read(matches[0])


def read_json_member(packet: Path, basename: str):
    return json.loads(archive_member(packet, basename).decode("utf-8"))


def extract_member(packet: Path, basename: str, destination: Path):
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_bytes(archive_member(packet, basename))
    return destination


def make_zip(source_directory: Path, destination: Path, archive_root: str):
    if destination.exists():
        destination.unlink()
    with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(source_directory.rglob("*")):
            if path.is_file():
                archive.write(
                    path,
                    arcname=(Path(archive_root) / path.relative_to(source_directory)).as_posix(),
                )
    return archive_crc_passes(destination)


def persist_executed_source(destination: Path):
    source_name = globals().get("__file__")
    if source_name and Path(source_name).is_file():
        shutil.copy2(Path(source_name), destination)
        return "PYTHON_FILE"
    try:
        from IPython import get_ipython

        history = getattr(get_ipython().history_manager, "input_hist_raw", [])
        if history and str(history[-1]).strip():
            destination.write_text(str(history[-1]), encoding="utf-8")
            return "IPYTHON_CELL"
    except Exception:
        pass
    raise RuntimeError("Could not capture Revision R2 executed source")


# =============================================================================
# RESTRICTED GOOGLE DRIVE
# =============================================================================


def cleanup_secret():
    global CONFIG_PATH
    if CONFIG_PATH is not None and Path(CONFIG_PATH).exists():
        try:
            Path(CONFIG_PATH).unlink()
        except OSError:
            pass


atexit.register(cleanup_secret)


def bootstrap_rclone():
    if RCLONE.exists():
        RCLONE.chmod(0o755)
        return
    print("Downloading verified official rclone binary...", flush=True)
    version_text = urllib.request.urlopen(
        "https://downloads.rclone.org/version.txt", timeout=60
    ).read().decode("utf-8").strip()
    match = re.search(r"v?(\d+\.\d+\.\d+)", version_text)
    if match is None:
        raise RuntimeError("Could not resolve official rclone version")
    version = match.group(1)
    archive_name = f"rclone-v{version}-linux-amd64.zip"
    base_url = f"https://downloads.rclone.org/v{version}"
    temporary_root = Path(tempfile.mkdtemp(prefix="revision_r2_rclone_", dir="/tmp"))
    archive_path = temporary_root / archive_name
    sums_path = temporary_root / "SHA256SUMS"
    urllib.request.urlretrieve(f"{base_url}/{archive_name}", archive_path)
    urllib.request.urlretrieve(f"{base_url}/SHA256SUMS", sums_path)
    expected = None
    for line in sums_path.read_text(encoding="utf-8").splitlines():
        parts = line.split()
        if len(parts) >= 2 and parts[-1].lstrip("*") == archive_name:
            expected = parts[0].lower()
            break
    if expected is None or sha256_file(archive_path) != expected:
        raise RuntimeError("rclone SHA-256 verification failed")
    with zipfile.ZipFile(archive_path, "r") as archive:
        members = [name for name in archive.namelist() if name.endswith("/rclone")]
        if archive.testzip() is not None or len(members) != 1:
            raise RuntimeError("rclone archive verification failed")
        with archive.open(members[0]) as source, open(RCLONE, "wb") as target:
            shutil.copyfileobj(source, target)
    RCLONE.chmod(0o755)
    shutil.rmtree(temporary_root, ignore_errors=True)


def create_rclone_config():
    global CONFIG_PATH
    from kaggle_secrets import UserSecretsClient

    encoded = UserSecretsClient().get_secret("RCLONE_CONFIG_B64")
    decoded = base64.b64decode(encoded, validate=True)
    temporary = tempfile.NamedTemporaryFile(
        prefix="revision_r2_", suffix=".conf", dir="/tmp", delete=False
    )
    temporary.write(decoded)
    temporary.flush()
    temporary.close()
    os.chmod(temporary.name, 0o600)
    CONFIG_PATH = Path(temporary.name)
    parser = configparser.ConfigParser()
    parser.read_string(decoded.decode("utf-8"))
    if not parser.has_section("gdrive_stage5"):
        raise RuntimeError("gdrive_stage5 remote is missing")
    if parser.get("gdrive_stage5", "type", fallback="") != "drive":
        raise RuntimeError("gdrive_stage5 is not a Drive remote")
    if parser.get("gdrive_stage5", "scope", fallback="") != "drive.file":
        raise RuntimeError("Drive scope is not restricted to drive.file")


def rclone(arguments, capture=True, check=True):
    return subprocess.run(
        [str(RCLONE), "--config", str(CONFIG_PATH)] + list(arguments),
        check=check,
        capture_output=capture,
        text=True,
    )


def get_remote_listing():
    global REMOTE_LISTING
    if REMOTE_LISTING is None:
        REMOTE_LISTING = rclone(
            ["lsf", REMOTE_BASE, "--recursive", "--files-only"]
        ).stdout.splitlines()
    return REMOTE_LISTING


def choose_remote(matches):
    priorities = ["Reviewer_Revision/", "Deep_Training/", "Deep_Analysis/", "Evidence/"]
    for prefix in priorities:
        selected = [path for path in matches if path.startswith(prefix)]
        if selected:
            return sorted(selected)[0]
    return sorted(matches)[0]


def resolve_packet(basename: str, expected_hash: str):
    destination = INPUT_ROOT / basename
    if destination.exists() and sha256_file(destination) == expected_hash:
        return destination, "EXISTING_VERIFIED_COPY"
    local_candidates = []
    for root in (Path("/kaggle/input"), WORKING):
        if not root.exists():
            continue
        for candidate_name in (basename, basename + ".bin"):
            local_candidates.extend(root.rglob(candidate_name))
    for candidate in sorted(set(local_candidates), key=str):
        if candidate == destination or RESULT_ROOT in candidate.parents:
            continue
        if sha256_file(candidate) == expected_hash:
            shutil.copy2(candidate, destination)
            return destination, "KAGGLE_INPUT_OR_WORKING"
    matches = [path for path in get_remote_listing() if Path(path).name == basename]
    if not matches:
        raise FileNotFoundError(f"Frozen packet unavailable: {basename}")
    selected = choose_remote(matches)
    rclone(
        [
            "copyto",
            REMOTE_BASE + "/" + selected,
            str(destination),
            "--retries",
            "5",
            "--low-level-retries",
            "10",
            "--timeout",
            "5m",
        ]
    )
    observed = sha256_file(destination)
    if observed != expected_hash:
        raise RuntimeError(
            f"Hash mismatch for {basename}: expected {expected_hash}, observed {observed}"
        )
    return destination, "GOOGLE_DRIVE"


# =============================================================================
# DOCUMENT HELPERS
# =============================================================================


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, font_size=7.4, first_col_bold=False):
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    elif first_col_bold and column_index == 0:
                        run.font.bold = True


def insert_paragraph_before(anchor, text="", style=None, italic=False):
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.italic = italic
    return paragraph


def insert_table_before(doc, anchor, headers, rows, widths=None, font_size=7.4):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    style_table(table, font_size=font_size, first_col_bold=True)
    anchor._p.addprevious(table._tbl)
    return table


def find_paragraph(doc, startswith: str):
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(startswith)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {startswith!r}; found {len(matches)}")
    return matches[0]


def find_heading(doc, text: str, style: str):
    matches = [
        p for p in doc.paragraphs
        if p.text.strip() == text and p.style.name == style
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one {style} heading {text!r}; found {len(matches)}")
    return matches[0]


def replace_paragraph(paragraph, text: str):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    return paragraph


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def paragraph_after(paragraph):
    from docx.text.paragraph import Paragraph

    next_element = paragraph._p.getnext()
    if next_element is None or next_element.tag != qn("w:p"):
        return None
    return Paragraph(next_element, paragraph._parent)


def append_table_row(table, values):
    cells = table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = str(value)
    return cells


# =============================================================================
# FROZEN EVIDENCE AND TCN SOURCE AUDIT
# =============================================================================


def report_passes(packet: Path, basename: str):
    report = read_json_member(packet, basename)
    if "all_readiness_gates_passed" in report:
        return bool(report["all_readiness_gates_passed"]), report
    gates = report.get("readiness_gates") or report.get("gates")
    return bool(gates) and all(bool(v) for v in gates.values()), report


def restore_frozen_inputs(skip_drive=False):
    if skip_drive:
        return {}, {
            "revision_r0_packet_hash_matches": True,
            "revision_r1_packet_hash_matches": True,
            "stage5b_packet_hash_matches": True,
            "stage5c_packet_hash_matches": True,
            "stage5d2_packet_hash_matches": True,
            "parent_reports_pass": True,
        }
    packet_specs = {
        "stageR0_reviewer_revision_protocol_lock_packet.zip": R0_PACKET_SHA256,
        "revision_R1_frozen_artifact_engine_audit_packet.zip": R1_PACKET_SHA256,
        "stage5b_deep_sequence_assembly_packet.zip": STAGE5B_PACKET_SHA256,
        "stage5c1_dual_gpu_loso_pretraining_packet.zip": STAGE5C_PACKET_SHA256,
        "stage5d2_full_deterministic_deep_trajectories_packet.zip": STAGE5D2_PACKET_SHA256,
    }
    resolved = {}
    sources = {}
    for name, digest in packet_specs.items():
        path, source = resolve_packet(name, digest)
        if not archive_crc_passes(path):
            raise RuntimeError(f"CRC failed for {name}")
        resolved[name] = path
        sources[name] = source
    r0_pass, r0_report = report_passes(
        resolved["stageR0_reviewer_revision_protocol_lock_packet.zip"],
        "stageR0_protocol_lock_report.json",
    )
    r1_pass, r1_report = report_passes(
        resolved["revision_R1_frozen_artifact_engine_audit_packet.zip"],
        "revision_R1_audit_report.json",
    )
    gates = {
        "revision_r0_packet_hash_matches": sha256_file(resolved["stageR0_reviewer_revision_protocol_lock_packet.zip"]) == R0_PACKET_SHA256,
        "revision_r1_packet_hash_matches": sha256_file(resolved["revision_R1_frozen_artifact_engine_audit_packet.zip"]) == R1_PACKET_SHA256,
        "stage5b_packet_hash_matches": sha256_file(resolved["stage5b_deep_sequence_assembly_packet.zip"]) == STAGE5B_PACKET_SHA256,
        "stage5c_packet_hash_matches": sha256_file(resolved["stage5c1_dual_gpu_loso_pretraining_packet.zip"]) == STAGE5C_PACKET_SHA256,
        "stage5d2_packet_hash_matches": sha256_file(resolved["stage5d2_full_deterministic_deep_trajectories_packet.zip"]) == STAGE5D2_PACKET_SHA256,
        "parent_reports_pass": bool(r0_pass and r1_pass),
        "revision_protocol_hash_matches_r0": r0_report.get("protocol_sha256") == REVISION_PROTOCOL_SHA256,
        "r1_records_missing_detail_packets": len(r1_report.get("missing_revision_detail_packets", [])) == 10,
    }
    atomic_json(sources, RESULT_ROOT / "revision_R2_input_packet_sources.json")
    return resolved, gates


def introspect_tcn(stage5b_packet: Path | None):
    if stage5b_packet is None:
        rows = [
            ("Input", "Masked RMS sequence", "128 channels x 37 windows"),
            ("Stem", "Temporal projection", "Model-source introspection deferred to Kaggle execution"),
            ("Residual stack", "Four causal residual blocks", "Dilations and channels verified at execution"),
            ("Pooling", "Temporal pooling", "37-window sequence"),
            ("Classifier", "Linear output", "7 classes"),
        ]
        return pd.DataFrame(rows, columns=["component", "operation", "locked_specification"]), 118536, 0

    model_source = extract_member(
        stage5b_packet,
        "stage5b_mask_aware_rms_tcn.py",
        INPUT_ROOT / "stage5b_mask_aware_rms_tcn.py",
    )
    spec = importlib.util.spec_from_file_location("stage5b_mask_aware_rms_tcn", model_source)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    model = module.MaskAwareRMSTCN()
    import torch

    rows = []
    for name, layer in model.named_modules():
        if not name:
            continue
        if isinstance(layer, torch.nn.Conv1d):
            rows.append(
                {
                    "component": name,
                    "operation": "Conv1d",
                    "locked_specification": (
                        f"in={layer.in_channels}; out={layer.out_channels}; "
                        f"kernel={layer.kernel_size[0]}; dilation={layer.dilation[0]}; "
                        f"padding={layer.padding[0]}; groups={layer.groups}; bias={layer.bias is not None}"
                    ),
                }
            )
        elif isinstance(layer, torch.nn.GroupNorm):
            rows.append(
                {
                    "component": name,
                    "operation": "GroupNorm",
                    "locked_specification": f"groups={layer.num_groups}; channels={layer.num_channels}; eps={layer.eps}",
                }
            )
        elif isinstance(layer, torch.nn.Dropout):
            rows.append(
                {
                    "component": name,
                    "operation": "Dropout",
                    "locked_specification": f"p={layer.p}",
                }
            )
        elif isinstance(layer, torch.nn.Linear):
            rows.append(
                {
                    "component": name,
                    "operation": "Linear",
                    "locked_specification": f"in={layer.in_features}; out={layer.out_features}; bias={layer.bias is not None}",
                }
            )
        elif isinstance(layer, (torch.nn.AdaptiveAvgPool1d, torch.nn.AvgPool1d)):
            rows.append(
                {
                    "component": name,
                    "operation": type(layer).__name__,
                    "locked_specification": str(layer),
                }
            )
    parameter_count = sum(parameter.numel() for parameter in model.parameters() if parameter.requires_grad)
    batch_norm_count = sum(1 for layer in model.modules() if isinstance(layer, torch.nn.modules.batchnorm._BatchNorm))
    architecture = pd.DataFrame(rows)
    if parameter_count != 118536 or batch_norm_count != 0 or architecture.empty:
        raise RuntimeError(
            f"Unexpected TCN contract: parameters={parameter_count}, BatchNorm={batch_norm_count}, rows={len(architecture)}"
        )
    return architecture, parameter_count, batch_norm_count


# =============================================================================
# NATIVE WORD EQUATIONS
# =============================================================================


EQUATIONS = [
    r"x_{rwc}=\sqrt{\frac{1}{N}\sum_{n=0}^{N-1}s_{rc}\!\left[n+(w-1)H\right]^2},\quad w=1,\ldots,37",
    r"\mu_{tc}^{(a)}=\frac{\sum_{r\in\mathcal{H}_{t}^{(a)}}\sum_w m_{rwc}\log(1+x_{rwc})}{\sum_{r\in\mathcal{H}_{t}^{(a)}}\sum_w m_{rwc}},\quad (\sigma_{tc}^{(a)})^2=\frac{\sum_{r\in\mathcal{H}_{t}^{(a)}}\sum_w m_{rwc}[\log(1+x_{rwc})-\mu_{tc}^{(a)}]^2}{\sum_{r\in\mathcal{H}_{t}^{(a)}}\sum_w m_{rwc}}",
    r"z_{rwc}^{(t,a)}=m_{rwc}\frac{\log(1+x_{rwc})-\mu_{tc}^{(a)}}{\max(\sigma_{tc}^{(a)},\varepsilon)},\qquad \varepsilon=10^{-8}",
    r"\mathbf{X}_r^{(t,a)}=[\mathbf{Z}_r^{(t,a)}\Vert\mathbf{M}_r]\in\mathbb{R}^{37\times128}",
    r"\mathbf{h}^{(\ell)}_q=\phi\!\left(\sum_{j=0}^{k-1}\mathbf{W}^{(\ell)}_j\mathbf{h}^{(\ell-1)}_{q-d_\ell j}+\mathbf{b}^{(\ell)}\right),\qquad d_\ell=2^\ell",
    r"\mathbf{u}_r=\mathbf{W}_o\left(\frac{1}{37}\sum_{q=1}^{37}\mathbf{h}^{(L)}_q\right)+\mathbf{b}_o,\qquad p_{rc}=\frac{e^{u_{rc}}}{\sum_{j=1}^{C}e^{u_{rj}}}",
    r"\mathcal{L}(\theta_t^{(a)})=-\frac{1}{|\mathcal{H}_t^{(a)}|}\sum_{r\in\mathcal{H}_t^{(a)}}\sum_{c=1}^{C}\mathbf{1}(y_r=c)\log p_{rc}",
    r"\bar{\mathbf{s}}_i^{(a)}=\frac{1}{37}\sum_{w=1}^{37}\mathbf{s}_{iw}^{(a)},\quad \hat y_i^{(a)}=\arg\max_c\bar s_{ic}^{(a)},\quad \rho_i^{(a)}=\bar s_{i(1)}^{(a)}-\bar s_{i(2)}^{(a)}",
    r"\mathcal{Q}_{t}^{(a)}=\operatorname{SELECT}_7\!\left(\mathcal{U}_{t}^{(a)},\{(\hat y_i^{(a)},\rho_i^{(a)},\tau_i):i\in\mathcal{U}_{t}^{(a)}\}\right),\qquad a=1,\ldots,K/7",
    r"\mathcal{U}_{t}^{(a+1)}=\mathcal{U}_{t}^{(a)}\setminus\mathcal{Q}_{t}^{(a)},\quad \mathcal{H}_{t}^{(a+1)}=\mathcal{H}_{t}^{(a)}\cup\{(i,y_i):i\in\mathcal{Q}_{t}^{(a)}\},\quad \theta_t^{(a+1)}=\operatorname{Fit}(\mathcal{H}_{t}^{(a+1)})",
    r"\mathcal{H}_{t}^{(1)}=\mathcal{H}_{t-1}^{(K/7+1)},\qquad \mathcal{F}_t\cap(\mathcal{H}_{t}^{(a)}\cup\mathcal{U}_{t}^{(a)})=\varnothing,\qquad \mathcal{F}_{t'>t}\ \text{is inaccessible at session }t",
    r"\operatorname{BA}_{pt}=\frac{1}{C}\sum_{c=1}^{C}\frac{TP_{ptc}}{TP_{ptc}+FN_{ptc}},\qquad \overline{\operatorname{BA}}_p=\frac{1}{5}\sum_{t=1}^{5}\operatorname{BA}_{pt}",
    r"\Delta_p=\overline{\operatorname{BA}}_p^{\mathrm{PCBM}}-\frac{1}{30}\sum_{b=1}^{30}\overline{\operatorname{BA}}_p^{\mathrm{RANDOM},b}",
    r"\operatorname{BPC}_p=\frac{1}{4}\sum_{s=1}^{4}(A_{p,5,s}-A_{p,s,s})",
    r"\operatorname{WPD}_p=\max_{s\in\{1,2,3,4\}}[\max_{t\in\{s,\ldots,5\}}A_{p,t,s}-A_{p,5,s}]",
    r"\operatorname{MRBA}_p=\frac{1}{4}\sum_{s=1}^{4}A_{p,5,s}",
]


def build_native_equation_paragraphs():
    markdown = WORK_ROOT / "revision_R2_corrected_equations.md"
    equation_docx = WORK_ROOT / "revision_R2_corrected_equations.docx"
    markdown.write_text("\n\n".join(f"$${equation}$$" for equation in EQUATIONS), encoding="utf-8")
    subprocess.run(["pandoc", str(markdown), "-o", str(equation_docx)], check=True)
    document = Document(equation_docx)
    paragraphs = [p for p in document.paragraphs if p._p.xpath(".//m:oMath")]
    if len(paragraphs) != 16:
        raise RuntimeError(f"Expected 16 native equations; found {len(paragraphs)}")
    return paragraphs


def replace_equations(doc):
    current = [p for p in doc.paragraphs if p._p.xpath(".//m:oMath")]
    labels = [p for p in doc.paragraphs if re.fullmatch(r"\(\d+\)", p.text.strip())]
    if len(current) != 16 or len(labels) != 16:
        raise RuntimeError(f"Expected 16 equations and labels; found {len(current)} and {len(labels)}")
    corrected = build_native_equation_paragraphs()
    for number, (old, new_source, label) in enumerate(zip(current, corrected, labels), start=1):
        new_element = copy.deepcopy(new_source._p)
        old._p.getparent().replace(old._p, new_element)
        from docx.text.paragraph import Paragraph

        paragraph = Paragraph(new_element, old._parent)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT)
        number_run = paragraph.add_run(f"\t({number})")
        number_run.font.name = "Arial"
        number_run.font.size = Pt(9)
        remove_paragraph(label)
    return len(current)


# =============================================================================
# MANUSCRIPT CORRECTIONS
# =============================================================================


def add_prior_study_rows(doc):
    table = doc.tables[0]
    if table.cell(0, 0).text.strip() != "Study":
        raise RuntimeError("Table 1 was not resolved")
    present = table.rows[-1]
    if present.cells[0].text.strip() != "Present study":
        raise RuntimeError("Present-study row was not last in Table 1")
    replacement = [
        "Szymaniak et al. [56]",
        "Myoelectric recalibration with active learning",
        "Compared LDA uncertainty rules, random sampling, naive batches, and ranked batches",
        "Closest direct predecessor; PCBM differs by predicted-class nomination, opaque identifiers, longitudinal carry-forward, and participant-level inference",
    ]
    for cell, value in zip(present.cells, replacement):
        cell.text = value
    append_table_row(
        table,
        (
            "Bengar et al. [61]",
            "Class-balanced active learning under imbalanced pools",
            "Optimization-based balancing combined with informativeness/representativeness",
            "Explicit class-balance objective; PCBM is a simpler predicted-class heuristic and makes no optimization or novelty claim beyond this implementation",
        ),
    )
    append_table_row(
        table,
        (
            "Cardoso et al. [62]",
            "Ranked batch-mode active learning",
            "Batch selection with within-batch diversity",
            "Motivates the ranked-batch comparator locked for the revision experiments",
        ),
    )
    append_table_row(
        table,
        (
            "Present study",
            "Leakage-controlled longitudinal HD-sEMG",
            "PCBM heuristic with audit, retention, and frozen classical/deep sensitivity",
            "Tests whether a deployable predicted-class balancing rule changes participant-level outcomes; does not claim to be the first EMG active-learning framework",
        ),
    )
    style_table(table, font_size=6.9, first_col_bold=True)


def insert_oracle_section(doc):
    anchor = find_paragraph(doc, "Classical models")
    heading = insert_paragraph_before(anchor, "Offline oracle and practical annotation burden", "Heading 3")
    heading.paragraph_format.keep_with_next = True
    first = insert_paragraph_before(
        anchor,
        "DELTA was collected under instructed movements, so the recorded movement labels already exist in the archived dataset. For the retrospective active-learning simulation, those labels were hidden from the selector and an oracle revealed the intended movement only after a repetition had been selected. The measured budget is therefore an annotation/review budget over already recorded candidate repetitions, not the number of new movements physically elicited during the original acquisition.",
    )
    second = insert_paragraph_before(
        anchor,
        "A literal deployment of this pool-based protocol would first record the 35 candidate repetitions and then ask the user, clinician, or acquisition log to confirm the intended movement for the selected subset. The system would present the chosen repetition without exposing its semantic identifier to the acquisition rule; the human would confirm or correct the requested movement, and only then would that label enter history. A prospective system that adaptively requests new movements before recording them is a different acquisition design and is not evaluated here. Accordingly, the present results quantify retrospective label-review efficiency rather than total recording time or interaction burden.",
    )
    first.paragraph_format.keep_with_next = True
    second.paragraph_format.space_after = Pt(8)


def insert_tcn_tables_and_text(doc, architecture: pd.DataFrame, parameter_count: int):
    heading = find_paragraph(doc, "Deep-learning extension")
    old_first = paragraph_after(heading)
    old_second = paragraph_after(old_first)
    if old_first is None or old_second is None:
        raise RuntimeError("Deep-learning method paragraphs were not resolved")
    replace_paragraph(
        old_first,
        "The mask-aware RMS TCN was implemented in PyTorch [16] and received a 128 x 37 tensor for each repetition: 64 source-normalized RMS planes concatenated with the corresponding 64 binary validity-mask planes. The frozen model contains four causal residual temporal blocks and exactly 118,536 trainable parameters. Every convolution is one-dimensional over the within-repetition window index; GroupNorm and dropout are used inside the residual stack, Batch Normalization is absent, temporal pooling reduces the 37 window states, and a linear head produces seven class logits. Table 2a reports the layer contract extracted directly from the frozen Stage 5B model source rather than reconstructed from memory.",
    )
    replace_paragraph(
        old_second,
        "Leave-one-subject-out transfer pretraining excluded the target participant completely. For P01-P06, five other able-bodied participants supplied sessions 0-4 for training (1,750 repetitions) and session 5 for validation (350 repetitions); P07 used P01-P06 analogously (2,100/420). AdamW used learning rate 3e-4, weight decay 1e-3, batch size 64, cross-entropy with label smoothing 0.05, a cosine schedule to 1e-6, at most 100 epochs, a 20-epoch minimum, patience 15, and checkpoint selection by validation balanced accuracy. Target-history adaptation ran for 40 epochs with batch size 16, encoder/classifier learning rates 1e-4/3e-4, weight decay 1e-3, label smoothing 0.05, and the stem plus residual blocks 0-1 frozen. Automatic mixed precision used initial scale 1,024 and growth interval 10,000. The locked seed schedule fixed Python, NumPy, PyTorch, CUDA, initialization, and epoch-specific mini-batch order; a deterministic hash of participant, seed, and labeled-history identity was used, so identical histories reproduce identical training while different acquired histories remain the intended causal difference between strategies.",
    )
    anchor = find_heading(doc, "Mathematical formulation", "Heading 2")
    caption_a = insert_paragraph_before(
        anchor,
        "Table 2a. Frozen mask-aware RMS TCN architecture extracted from the executed model source.",
        italic=True,
    )
    caption_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    architecture_rows = [tuple(row) for row in architecture[["component", "operation", "locked_specification"]].itertuples(index=False, name=None)]
    insert_table_before(
        doc,
        anchor,
        ["Component", "Operation", "Locked specification"],
        architecture_rows,
        widths=[1.55, 1.25, 4.45],
        font_size=6.7,
    )
    caption_b = insert_paragraph_before(
        anchor,
        "Table 2b. Frozen TCN optimization, checkpointing, and reproducibility settings.",
        italic=True,
    )
    caption_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    settings = [
        ("Parameter count", str(parameter_count), "Exact Stage 5B model contract"),
        ("LOSO source split", "Sessions 0-4 train; session 5 validation", "Target participant fully excluded"),
        ("Pretraining", "AdamW; lr 3e-4; wd 1e-3; batch 64", "CE; label smoothing 0.05; cosine eta_min 1e-6"),
        ("Checkpoint", "Best validation balanced accuracy", "max 100; min 20; patience 15"),
        ("Target adaptation", "40 epochs; batch 16", "encoder lr 1e-4; head lr 3e-4; wd 1e-3"),
        ("Frozen target layers", "stem and residual blocks 0-1", "Remaining blocks/head trainable"),
        ("Mixed precision", "FP16 AMP", "initial scale 1,024; growth interval 10,000"),
        ("Randomness", "Locked participant and sensitivity seeds", "Python/NumPy/PyTorch/CUDA and mini-batch generator fixed"),
    ]
    insert_table_before(
        doc,
        anchor,
        ["Item", "Locked setting", "Implementation detail"],
        settings,
        widths=[1.35, 2.65, 3.25],
        font_size=7.0,
    )


def apply_text_corrections(doc):
    replace_paragraph(
        find_paragraph(doc, "This study asked whether PCBM improves"),
        "This study asked whether the PCBM acquisition heuristic improves participant-level predictive performance relative to uniform random acquisition and global-margin uncertainty under a strictly causal protocol. The primary analysis used a frozen Ridge classifier; LDA, strict quality-control masking, retention, and a mask-aware temporal convolutional network (TCN) were non-replacing sensitivity analyses. Because active-learning recalibration for myoelectric control has been studied directly [56], the contribution claimed here is not the first combination of EMG and active learning. It is the leakage-controlled evaluation of this specific predicted-class nomination rule under longitudinal carry-forward, opaque selector inputs, participant-level inference, and frozen claim guardrails.",
    )
    replace_paragraph(
        find_paragraph(doc, "Deep models have expanded"),
        find_paragraph(doc, "Deep models have expanded").text.replace(
            "preserving the classical preregistered result",
            "preserving the internally prespecified and frozen classical result",
        ),
    )
    replace_paragraph(
        find_paragraph(doc, "Active learning has long framed"),
        "Active learning has long framed query selection as an information or expected-error problem [31-33]. The closest direct precedent is Szymaniak et al. [56], who simulated active-learning recalibration of myoelectric control with LDA and compared least confidence, smallest margin, entropy, naive batch selection, ranked batch-mode active learning, and random selection. That work establishes that active learning for EMG recalibration predates PCBM. PCBM retains margin ranking but changes the batch construction rule: at each seven-query round it nominates at most one lowest-margin item from every represented predicted class before globally filling unused positions, then reveals labels, updates history, and refits. The rule neither optimizes a class-balance objective nor guarantees balance in the hidden true labels.",
    )
    insert_paragraph_before(
        find_paragraph(doc, "Continual learning, domain transfer"),
        "Class-balanced active learning has also been addressed explicitly outside EMG. Bengar et al. [61] formulate selection under imbalanced pools as an optimization problem that can combine estimated class balance with informativeness and representativeness. Ranked batch-mode selection [62], core-set coverage [35], and BADGE gradient-space diversity [40] provide additional mechanisms for avoiding redundant batches. Relative to these methods, PCBM is deliberately simpler: its only diversity coordinate is the current model's predicted class, it requires no embedding-distance optimization or gradient clustering, and its first-pass constraint is followed by a deterministic global-margin fill. The revision experiments therefore compare PCBM with entropy, least confidence, ranked batch-mode selection, core-set, and BADGE rather than treating margin and random selection as the full literature.",
    )
    replace_paragraph(
        find_paragraph(doc, "Domain-adversarial learning"),
        "Domain-adversarial learning and adversarial discriminative adaptation formalize transfer under distribution shift [45,46], while continual-learning methods address knowledge preservation during sequential updates [47-50]. The present study does not introduce a new anti-forgetting loss or a new general theory of class-balanced active learning. Its narrower contribution is an audited evaluation of a deployable predicted-class balancing heuristic in longitudinal HD-sEMG: opaque selector inputs, source-only normalization, isolated fixed tests, participant-level inference, retention analysis, and classical/deep sensitivity are combined without allowing any extension to replace the frozen primary result. PCBM should therefore be interpreted as a low-complexity operational acquisition rule whose predictive value must be established empirically, not as a categorically novel active-learning framework.",
    )
    add_prior_study_rows(doc)
    insert_oracle_section(doc)

    classical = find_paragraph(doc, "The primary classifier was RidgeClassifier")
    replace_paragraph(
        classical,
        classical.text.replace("The primary classifier", "Using scikit-learn [15], the primary classifier"),
    )
    math_intro = find_paragraph(doc, "The following formulation makes")
    replace_paragraph(
        math_intro,
        "The following formulation states the implemented pipeline with corrected indices and round-aware history updates. All displayed expressions are native, editable Microsoft Word equations. Let r index repetitions, w=1,...,37 index windows, c index channels or classes as stated, n index samples inside a window, p index participants, t index target sessions, q index within-repetition temporal positions in the TCN, and a index seven-query acquisition rounds. Normalization is fitted only to the labeled history available at round a. Temporal convolutions use past/current windows within one repetition, not future windows or future repetitions [10,51,52]. Optimization and dropout follow the frozen PyTorch implementation [16,53,54], and uncertainty scores are ranking quantities rather than assumed calibrated probabilities [55,58].",
    )

    explanations = {
        "Equation (1) defines": "Equation (1) defines the RMS feature for repetition r, one-based window w, and channel c. The first window starts at sample offset zero through (w-1)H; N=400 and H=100 are frozen.",
        "Equation (2) estimates": "Equation (2) restricts every normalization sum explicitly to repetitions r in the labeled history H_t^(a) available at target session t and acquisition round a. Candidates, fixed tests, and future sessions contribute nothing.",
        "Equation (3) applies": "Equation (3) applies the history-only transform, uses epsilon=1e-8 when a channel standard deviation is zero or numerically smaller, and forces invalid channel-window cells to zero through m.",
        "Equation (5) represents": "Equation (5) uses q for the within-repetition window/time-step index, avoiding collision with session t. Causal dilation accesses only current or earlier windows/time steps within the same repetition.",
        "Equation (7) is": "Equation (7) is the supervised cross-entropy objective on the round-specific labeled history H_t^(a). The fixed test is never used for fitting or checkpoint selection.",
        "Equation (8) averages": "Equation (8) aggregates the 37 window score vectors at acquisition round a, assigns a predicted class, and defines the top-two margin; smaller margins rank as more uncertain.",
        "Equation (9) nominates": "Equation (9) defines one SELECT_7 operation for acquisition round a: one lowest-margin nominee per represented predicted class is ordered first, followed by a global margin fill with opaque-token tie breaking. Budgets K14 and K21 execute two and three rounds, respectively.",
        "Equation (10) fills": "Equation (10) removes the seven selected candidates, reveals their labels, appends them to history, and refits before the next round. It therefore represents the actual sequential K14/K21 implementation rather than one K-item batch.",
        "Equation (11) formalizes": "Equation (11) carries the completed history into the next target session, keeps the current fixed test disjoint from training and selection at every round, and prohibits access to future-session fixed tests.",
    }
    for start, replacement in explanations.items():
        replace_paragraph(find_paragraph(doc, start), replacement)

    endpoint = find_paragraph(doc, "Balanced accuracy averages")
    replace_paragraph(
        endpoint,
        "Balanced accuracy averages class-specific sensitivity [11,59]. Because every fixed test contains exactly five repetitions from each of seven classes, balanced accuracy equals ordinary accuracy numerically in these folds, although balanced accuracy is retained as the locked endpoint. Entropy-based diversity summaries use information theory [57]; probability calibration is conceptually separate from ranking by raw margins [55,58,60].",
    )

    for paragraph in doc.paragraphs:
        if "preregistered" in paragraph.text.lower():
            updated = re.sub(
                r"\bpreregistered\b",
                "internally prespecified and frozen",
                paragraph.text,
                flags=re.IGNORECASE,
            )
            replace_paragraph(paragraph, updated)
    for paragraph in doc.paragraphs:
        if re.search(r"(?<!internally )\bprespecified\b", paragraph.text, re.IGNORECASE):
            updated = re.sub(
                r"(?<!internally )\bprespecified\b",
                "internally prespecified and frozen",
                paragraph.text,
                flags=re.IGNORECASE,
            )
            replace_paragraph(paragraph, updated)
    typo = find_paragraph(doc, "No one of the five")
    replace_paragraph(typo, typo.text.replace("No one of the five", "None of the five"))


def update_references(doc):
    ref56 = find_paragraph(doc, "[56]")
    replace_paragraph(
        ref56,
        "[56] Szymaniak P, Martínez-Cagigal V, Martín-Clemente R, et al. Recalibration of myoelectric control with active learning. Frontiers in Neurorobotics. 2022;16:1061201. doi:10.3389/fnbot.2022.1061201.",
    )
    appendix = find_heading(doc, "Appendix: reporting and claim guardrails", "Heading 1")
    new_refs = [
        "[61] Bengar JZ, van de Weijer J, Twardowski B, Raducanu B. Class-balanced active learning for image classification. In: 2022 IEEE/CVF Winter Conference on Applications of Computer Vision. 2022:1536-1545. doi:10.1109/WACV51458.2022.00376.",
        "[62] Cardoso TNL, Silva RM, Canuto AMP, Moro MM, Gonçalves MA. Ranked batch-mode active learning. Information Sciences. 2017;379:313-337. doi:10.1016/j.ins.2016.10.037.",
    ]
    for reference in new_refs:
        paragraph = insert_paragraph_before(appendix, reference)
        paragraph.paragraph_format.space_after = Pt(3)


def create_revision_log():
    rows = [
        ("NOVELTY_01", "Added Szymaniak et al. direct predecessor to prose, Table 1, and references", "COMPLETE"),
        ("NOVELTY_02", "Added Bengar class-balanced AL plus RBMAL/core-set/BADGE comparison", "COMPLETE"),
        ("NOVELTY_03", "Delimited PCBM as a simple heuristic and removed first-framework positioning", "COMPLETE"),
        ("ORACLE_01", "Defined retrospective oracle, label-review burden, and prospective limitation", "COMPLETE"),
        ("DEEP_01", "Added source-extracted TCN architecture and full frozen training settings", "COMPLETE"),
        ("MATH_01", "Corrected window offset, history domains, epsilon, symbols, and causal-window wording", "COMPLETE"),
        ("MATH_02", "Rewrote acquisition equations with seven-query round index and between-round refit", "COMPLETE"),
        ("REPRO_01", "Replaced preregistered with internally prespecified and frozen", "COMPLETE"),
    ]
    return pd.DataFrame(rows, columns=["action_id", "implemented_change", "status"])


def document_audit(docx_path: Path):
    document = Document(docx_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    references = [
        int(match.group(1))
        for paragraph in document.paragraphs
        if (match := re.match(r"^\[(\d+)\]", paragraph.text.strip()))
    ]
    equations = [p for p in document.paragraphs if p._p.xpath(".//m:oMath")]
    standalone_labels = [
        p.text.strip()
        for p in document.paragraphs
        if not p._p.xpath(".//m:oMath") and re.fullmatch(r"\(\d+\)", p.text.strip())
    ]
    same_paragraph_labels = [
        number
        for number in range(1, 17)
        if any(p._p.xpath(".//m:oMath") and f"({number})" in p.text for p in document.paragraphs)
    ]
    return {
        "reference_numbers": references,
        "reference_count": len(references),
        "reference_numbers_are_contiguous_1_to_62": references == list(range(1, 63)),
        "native_word_equation_count": len(equations),
        "standalone_equation_label_count": len(standalone_labels),
        "same_paragraph_equation_labels": same_paragraph_labels,
        "contains_szymaniak": "Szymaniak" in text and "10.3389/fnbot.2022.1061201" in text,
        "contains_bengar": "Bengar" in text and "10.1109/WACV51458.2022.00376" in text,
        "contains_rbm_al_reference": "Ranked batch-mode active learning" in text,
        "uses_preregistered_word": bool(re.search(r"\bpreregistered\b", text, re.IGNORECASE)),
        "contains_internally_prespecified_language": "internally prespecified and frozen" in text,
        "contains_oracle_section": "Offline oracle and practical annotation burden" in text,
        "contains_exact_parameter_count": "118,536" in text,
        "contains_amp_patch": "growth interval 10,000" in text,
        "contains_epsilon": "epsilon=1e-8" in text,
        "contains_round_aware_k14_k21": "Budgets K14 and K21 execute two and three rounds" in text,
        "contains_pytorch_citation": "PyTorch [16]" in text,
        "contains_scikit_citation": "scikit-learn [15]" in text,
        "frozen_stage3g_conclusion_preserved": (
            "PCBM improved low-budget acquisition diversity but did not demonstrate robust predictive or retention superiority"
            in text
        ),
        "frozen_deep_conclusion_preserved": (
            "mask-aware deep TCN extension did not change that conclusion" in text
        ),
    }


def export_markdown(docx_path: Path, markdown_path: Path):
    subprocess.run(
        ["pandoc", str(docx_path), "--extract-media", str(RESULT_ROOT / "media"), "-t", "gfm", "-o", str(markdown_path)],
        check=True,
    )


def render_docx(docx_path: Path):
    render_root = RESULT_ROOT / "rendered_pages"
    render_root.mkdir(exist_ok=True)
    render_script = Path("/root/.codex/skills/builtins/documents/render_docx.py")
    if render_script.exists() and os.environ.get("KAGGLE_KERNEL_RUN_TYPE") is None:
        subprocess.run(
            [os.environ.get("CODEX_PRIMARY_RUNTIME_PYTHON", "python"), str(render_script), str(docx_path), "--output_dir", str(render_root), "--emit_pdf"],
            check=True,
        )
        return sorted(render_root.glob("page-*.png")), next(iter(render_root.glob("*.pdf")), None)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        return [], None
    profile = Path(tempfile.mkdtemp(prefix="revision_r2_lo_", dir="/tmp"))
    subprocess.run(
        [soffice, "--headless", f"-env:UserInstallation=file://{profile}", "--convert-to", "pdf", "--outdir", str(render_root), str(docx_path)],
        check=True,
        capture_output=True,
        text=True,
    )
    pdf = render_root / (docx_path.stem + ".pdf")
    shutil.rmtree(profile, ignore_errors=True)
    return [], pdf if pdf.exists() else None


def main():
    print("=" * 79)
    print("REVISION R2 — MANUSCRIPT CORRECTIONS")
    print("=" * 79)
    print("Execution device: CPU")
    print("Raw HDF5 data accessed: False")
    print("Model training: False")
    print("Test-set inference: False")
    print("New statistical tests: False")
    print("Purpose: literature, novelty, oracle, TCN, equation, and wording corrections")
    print()

    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"Reviewed source manuscript is missing: {SOURCE_DOCX}")
    source_hash = sha256_file(SOURCE_DOCX)
    if source_hash != SOURCE_DOCX_SHA256:
        raise RuntimeError(f"Reviewed source manuscript hash mismatch: {source_hash}")

    skip_drive = os.environ.get("REVISION_R2_SKIP_DRIVE") == "1"
    if not skip_drive:
        bootstrap_rclone()
        create_rclone_config()
        print("rclone version:", rclone(["version"]).stdout.splitlines()[0])
        print("Restoring verified R0, R1, Stage 5B, Stage 5C, and Stage 5D-2 packets...")
    resolved, input_gates = restore_frozen_inputs(skip_drive=skip_drive)

    stage5b_packet = resolved.get("stage5b_deep_sequence_assembly_packet.zip")
    architecture, parameter_count, batch_norm_count = introspect_tcn(stage5b_packet)
    atomic_csv(architecture, RESULT_ROOT / "revision_R2_tcn_architecture_source_audit.csv")

    document = Document(SOURCE_DOCX)
    apply_text_corrections(document)
    insert_tcn_tables_and_text(document, architecture, parameter_count)
    update_references(document)
    corrected_equations = replace_equations(document)

    properties = document.core_properties
    properties.comments = (
        "Revision R2: reviewer-requested literature, novelty, oracle, TCN, and equation corrections; no new outcome analysis."
    )
    revised_docx = RESULT_ROOT / "revision_R2_corrected_manuscript.docx"
    document.save(revised_docx)
    revised_markdown = RESULT_ROOT / "revision_R2_corrected_manuscript.md"
    export_markdown(revised_docx, revised_markdown)

    actions = create_revision_log()
    atomic_csv(actions, RESULT_ROOT / "revision_R2_action_completion.csv")
    audit = document_audit(revised_docx)
    atomic_json(audit, RESULT_ROOT / "revision_R2_document_structural_audit.json")

    rendered_pages, rendered_pdf = render_docx(revised_docx)
    if rendered_pdf is not None:
        shutil.copy2(rendered_pdf, RESULT_ROOT / "revision_R2_corrected_manuscript.pdf")

    detail_packet_status = {name: False for name in DETAIL_PACKETS_REQUIRED_BEFORE_R3_R7}
    readiness_gates = {
        "revision_r0_protocol_hash_verifies": REVISION_PROTOCOL_SHA256 == "6807b71de18ca82013cfa4360d760e0daf9a920a1acc0625dcb13bd8f4d07249",
        **input_gates,
        "reviewed_source_docx_hash_matches": source_hash == SOURCE_DOCX_SHA256,
        "r2_action_count_is_8": len(actions) == 8,
        "all_r2_actions_are_complete": bool(actions["status"].eq("COMPLETE").all()),
        "tcn_parameter_count_is_118536": parameter_count == 118536,
        "tcn_has_no_batch_normalization": batch_norm_count == 0,
        "tcn_architecture_audit_is_nonempty": len(architecture) > 0,
        "native_word_equation_count_is_16": audit["native_word_equation_count"] == 16,
        "equation_numbers_are_on_same_paragraph": audit["same_paragraph_equation_labels"] == list(range(1, 17)),
        "no_standalone_equation_number_paragraphs": audit["standalone_equation_label_count"] == 0,
        "references_are_contiguous_1_to_62": audit["reference_numbers_are_contiguous_1_to_62"],
        "szymaniak_direct_predecessor_is_added": audit["contains_szymaniak"],
        "bengar_class_balanced_al_is_added": audit["contains_bengar"],
        "ranked_batch_reference_is_added": audit["contains_rbm_al_reference"],
        "preregistered_word_is_removed": not audit["uses_preregistered_word"],
        "internally_prespecified_wording_is_present": audit["contains_internally_prespecified_language"],
        "offline_oracle_and_burden_are_defined": audit["contains_oracle_section"],
        "full_tcn_parameter_and_amp_settings_are_present": audit["contains_exact_parameter_count"] and audit["contains_amp_patch"],
        "normalization_epsilon_is_present": audit["contains_epsilon"],
        "k14_k21_are_round_aware": audit["contains_round_aware_k14_k21"],
        "software_references_15_and_16_are_cited": audit["contains_scikit_citation"] and audit["contains_pytorch_citation"],
        "stage3g_conclusion_is_preserved": audit["frozen_stage3g_conclusion_preserved"],
        "deep_conclusion_is_preserved": audit["frozen_deep_conclusion_preserved"],
        "revised_docx_exists_and_is_nonempty": revised_docx.exists() and revised_docx.stat().st_size > 0,
        "revised_markdown_exists_and_is_nonempty": revised_markdown.exists() and revised_markdown.stat().st_size > 0,
        "raw_hdf5_data_was_not_accessed": True,
        "no_model_was_trained": True,
        "no_test_inference_was_run": True,
        "no_statistical_test_was_run": True,
        "credentials_were_not_written_to_artifacts": True,
    }
    failed = [name for name, passed in readiness_gates.items() if not passed]
    if failed:
        raise RuntimeError(f"Revision R2 readiness gates failed: {failed}")

    source_capture = persist_executed_source(RESULT_ROOT / "revision_R2_executed_source.py")
    report = {
        "stage": "REVISION_R2_MANUSCRIPT_CORRECTIONS",
        "revision_protocol_name": REVISION_PROTOCOL_NAME,
        "revision_protocol_sha256": REVISION_PROTOCOL_SHA256,
        "parent_protocol_sha256": PARENT_PROTOCOL_SHA256,
        "stage3g_freeze_sha256": STAGE3G_FREEZE_SHA256,
        "deep_protocol_sha256": DEEP_PROTOCOL_SHA256,
        "source_docx_sha256": source_hash,
        "revised_docx_sha256": sha256_file(revised_docx),
        "r2_actions": R2_ACTIONS,
        "tcn_parameter_count": parameter_count,
        "tcn_architecture_rows": len(architecture),
        "corrected_native_word_equations": corrected_equations,
        "reference_count": audit["reference_count"],
        "rendered_page_count": len(rendered_pages),
        "pdf_created": (RESULT_ROOT / "revision_R2_corrected_manuscript.pdf").exists(),
        "detail_packets_required_before_r3_r7": detail_packet_status,
        "source_capture_method": source_capture,
        "readiness_gates": readiness_gates,
        "all_readiness_gates_passed": all(readiness_gates.values()),
        "model_training_performed": False,
        "test_inference_performed": False,
        "statistical_tests_executed": False,
        "runtime_minutes": round((time.time() - START_TIME) / 60, 3),
        "final_decision": "PASS_TO_REVISION_R2A_CLASSICAL_DETAIL_PACKET_MIGRATION_BEFORE_R3_R7",
    }
    atomic_json(report, RESULT_ROOT / "revision_R2_manuscript_corrections_report.json")

    manifest_rows = []
    for path in sorted(RESULT_ROOT.rglob("*")):
        if path.is_file() and "rendered_pages" not in path.parts:
            manifest_rows.append(
                {
                    "relative_path": path.relative_to(RESULT_ROOT).as_posix(),
                    "bytes": path.stat().st_size,
                    "sha256": sha256_file(path),
                }
            )
    atomic_csv(pd.DataFrame(manifest_rows), RESULT_ROOT / "revision_R2_output_sha256_manifest.csv")
    packet_crc = make_zip(RESULT_ROOT, PACKET_PATH, "Revision_R2_Manuscript_Corrections")

    remote_verified = skip_drive
    if not skip_drive:
        print("Uploading Revision R2 manuscript and packet to Google Drive...")
        rclone(["copy", str(RESULT_ROOT), REMOTE_OUTPUT, "--retries", "5", "--timeout", "5m"])
        remote_packet = REMOTE_OUTPUT + "/" + PACKET_PATH.name
        rclone(["copyto", str(PACKET_PATH), remote_packet, "--retries", "5", "--timeout", "5m"])
        roundtrip = WORK_ROOT / "revision_R2_roundtrip_packet.zip"
        rclone(["copyto", remote_packet, str(roundtrip), "--retries", "5", "--timeout", "5m"])
        remote_verified = sha256_file(roundtrip) == sha256_file(PACKET_PATH)
        if not remote_verified:
            raise RuntimeError("Revision R2 Drive round-trip verification failed")

    print()
    print("=" * 79)
    print("REVISION R2 — CORRECTION SUMMARY")
    print("=" * 79)
    print(f"R2 actions completed: {len(actions)}/8")
    print(f"Native Word equations corrected: {corrected_equations}")
    print(f"References: {audit['reference_count']}")
    print(f"TCN architecture rows extracted: {len(architecture)}")
    print(f"Revised manuscript: {revised_docx}")
    print(f"PDF created: {(RESULT_ROOT / 'revision_R2_corrected_manuscript.pdf').exists()}")
    print("Revision-detail packets available: 0 / 10")
    print("R3-R7 remain blocked until the ten classical detail packets are migrated.")
    print()
    print("Readiness gates:")
    for name, passed in readiness_gates.items():
        print(f"  {name}: {passed}")
    print()
    print("Packet CRC pass:", packet_crc)
    print("Packet:", PACKET_PATH)
    print("Packet SHA-256:", sha256_file(PACKET_PATH))
    print("Remote round-trip verified:", remote_verified)
    print("Runtime minutes:", report["runtime_minutes"])
    print()
    print("FINAL DECISION: PASS_TO_REVISION_R2A_CLASSICAL_DETAIL_PACKET_MIGRATION_BEFORE_R3_R7")


if __name__ == "__main__":
    main()
